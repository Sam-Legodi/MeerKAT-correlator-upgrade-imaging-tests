#!/usr/bin/env python3
"""
Flux catalogue comparison CLI.

This script mirrors the `Flux_analysis.ipynb` workflow: it ingests cross-matched
PyBDSF catalogues, generates flux comparison plots (with robust linear fits and
inlier/outlier identification), assembles summary statistics, and writes a DOCX
report alongside the figures.
"""

from __future__ import annotations

import argparse
import builtins
import glob
import os
from typing import Dict, Iterable, List, Optional, Sequence, Tuple

import matplotlib

matplotlib.use("Agg")  # non-interactive backend for batch runs
import matplotlib.pyplot as plt  # noqa: E402
import numpy as np  # noqa: E402
from astropy.table import Table  # noqa: E402
from docx import Document  # noqa: E402
from docx.shared import Inches  # noqa: E402
from docx.oxml import OxmlElement  # noqa: E402
from docx.oxml.ns import qn  # noqa: E402


# -------- shared helpers (safe column read + docx + figs) --------
def _finite_global_max(*arrs: Iterable[np.ndarray], default: float = 1.0) -> float:
    """Return the global finite max across arrays, or `default` if none are finite."""
    vals: List[float] = []
    for arr in arrs:
        a = np.asarray(arr, dtype=float)
        a = a[np.isfinite(a)]
        if a.size:
            vals.append(float(np.nanmax(a)))
    return float(np.nanmax(vals)) if vals else float(default)


def _ensure_dir(path: str) -> str:
    os.makedirs(path, exist_ok=True)
    return path


def _save_fig(path: str, *, tight: bool = True, dpi: int = 150) -> None:
    if tight:
        plt.tight_layout()
    plt.savefig(path, dpi=dpi)
    plt.close()


def _docx_add_caption(doc: Document, text: str) -> None:
    para = doc.add_paragraph(text)
    try:
        para.style = "Caption"
    except Exception:
        pass


def _safe_arr(values: Sequence) -> np.ndarray:
    """Convert an input column/sequence to a float numpy array, unwrapping Quantities/masked arrays."""
    if hasattr(values, "to_value"):
        values = values.to_value()
    arr = np.asarray(values)
    if np.ma.isMaskedArray(arr):
        arr = arr.filled(np.nan)
    return arr.astype(float, copy=False)


def _col_or_default(tab: Table, colname: str, *, like: Optional[np.ndarray] = None) -> np.ndarray:
    """Return named column as float ndarray if present, else NaNs with the same shape as `like`."""
    if colname in tab.colnames:
        return _safe_arr(tab[colname])
    like_arr = _safe_arr(like) if like is not None else np.array([], dtype=float)
    return np.full_like(like_arr, np.nan, dtype=float)


def _frac_diff(ref: np.ndarray, test: np.ndarray) -> np.ndarray:
    """(test - ref) / ref with divide-by-zero protection."""
    ref = np.asarray(ref, dtype=float)
    test = np.asarray(test, dtype=float)
    denom = np.where(np.abs(ref) > 0, ref, np.nan)
    return (test - ref) / denom


def _ratio(test: np.ndarray, ref: np.ndarray) -> np.ndarray:
    """test / ref with divide-by-zero protection."""
    ref = np.asarray(ref, dtype=float)
    test = np.asarray(test, dtype=float)
    denom = np.where(np.abs(ref) > 0, ref, np.nan)
    return test / denom


def _axis_inlier_mask(values: np.ndarray, max_z: float = 4.0) -> np.ndarray:
    """
    Return True for samples within `max_z` robust-sigma of the median.
    Helps cull extreme x/y points before regression so horizontal/vertical
    outliers cannot sneak into the fit.
    """
    arr = np.asarray(values, dtype=float).ravel()
    mask = np.zeros_like(arr, dtype=bool)
    if arr.size == 0:
        return mask
    finite = np.isfinite(arr)
    if not np.any(finite):
        return mask
    good = arr[finite]
    median = np.nanmedian(good)
    mad = 1.4826 * np.nanmedian(np.abs(good - median))
    if not np.isfinite(mad) or mad <= 1.0e-9:
        mad = np.nanstd(good)
    if not np.isfinite(mad) or mad <= 1.0e-9:
        mask[finite] = True
        return mask
    z = np.abs(good - median) / mad
    mask[finite] = z <= max_z
    return mask


def _fmt_stat(arr: np.ndarray, func, fmt: str = "{:.6g}", default: str = "nan") -> str:
    arr = np.asarray(arr, dtype=float)
    arr = arr[np.isfinite(arr)]
    if arr.size == 0:
        return default
    try:
        value = func(arr)
    except Exception:
        return default
    if value is None or not np.isfinite(value):
        return default
    return fmt.format(value)


def _summary_rows(
    label: str,
    frac: np.ndarray,
    ratio: Optional[np.ndarray] = None,
    gain: Optional[Tuple[float, str]] = None,
) -> List[List[str]]:
    """Build a summary table (as rows) for DOCX output."""
    frac = np.asarray(frac, dtype=float)
    frac_finite = frac[np.isfinite(frac)]
    rows: List[List[str]] = [["metric", f"{label}: value"]]
    rows.append(["N (finite frac)", str(frac_finite.size)])
    rows.append(["mean(frac)", _fmt_stat(frac_finite, np.nanmean)])
    rows.append(["std(frac)", _fmt_stat(frac_finite, np.nanstd)])
    rows.append(["median(frac)", _fmt_stat(frac_finite, np.nanmedian)])
    rows.append(["p16(frac)", _fmt_stat(frac_finite, lambda a: np.nanpercentile(a, 16))])
    rows.append(["p84(frac)", _fmt_stat(frac_finite, lambda a: np.nanpercentile(a, 84))])

    if ratio is not None:
        ratio = np.asarray(ratio, dtype=float)
        ratio_finite = ratio[np.isfinite(ratio)]
        rows.append(["N (finite ratio)", str(ratio_finite.size)])
        rows.append(["median(ratio)", _fmt_stat(ratio_finite, np.nanmedian)])
        rows.append(
            [
                "NMAD(ratio)",
                _fmt_stat(
                    ratio_finite,
                    lambda a: 1.4826 * np.nanmedian(np.abs(a - np.nanmedian(a))),
                ),
            ]
        )

    if gain is not None:
        g, note = gain
        rows.append(["gain (through-origin LS)", f"{g:.6g}"])
        rows.append(["gain notes", note])

    return rows

def _set_table_borders(table, width_pt: float = 1.0, color: str = "000000") -> None:
    """Ensure DOCX table borders are visible with the requested width."""
    size = max(int(round(width_pt * 8)), 1)  # Word units: 1/8th points
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.append(tblPr)
    borders = tblPr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tblPr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = qn(f"w:{edge}")
        element = borders.find(tag)
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), str(size))
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def _add_table(doc: Document, rows: List[List[str]], caption_text: str) -> None:
    table = doc.add_table(rows=1, cols=len(rows[0]))
    hdr = table.rows[0].cells
    for j, heading in enumerate(rows[0]):
        hdr[j].text = str(heading)
    for row_values in rows[1:]:
        row = table.add_row().cells
        for j, value in enumerate(row_values):
            row[j].text = str(value)
    _set_table_borders(table, width_pt=1.0)
    _docx_add_caption(doc, caption_text)


def _gain_slope_through_origin(
    ref: np.ndarray,
    test: np.ndarray,
    ref_err: Optional[np.ndarray] = None,
    test_err: Optional[np.ndarray] = None,
    max_iter: int = 3,
) -> Tuple[float, str]:
    """
    Estimate gain g in test ≈ g * ref using (optionally) weighted LS through the origin.
    If errors are given, iterate weights: w = 1/(σ_test^2 + (g * σ_ref)^2).
    """
    ref = np.asarray(ref, dtype=float)
    test = np.asarray(test, dtype=float)
    mask = np.isfinite(ref) & np.isfinite(test)
    if not np.any(mask):
        return (float("nan"), "no finite data")

    x = ref[mask]
    y = test[mask]

    if ref_err is not None and test_err is not None:
        re = np.asarray(ref_err, dtype=float)
        te = np.asarray(test_err, dtype=float)
        re = re[mask]
        te = te[mask]
        g = np.nansum(x * y) / np.nansum(x * x)
        if not np.isfinite(g):
            g = 1.0
        for _ in range(max_iter):
            var = te ** 2 + (g * re) ** 2
            w = 1.0 / np.where(var > 0, var, np.nan)
            g_new = np.nansum(w * x * y) / np.nansum(w * x * x)
            if not np.isfinite(g_new):
                break
            if np.allclose(g_new, g, rtol=1e-5, atol=0):
                g = g_new
                break
            g = g_new
        return (float(g), "weighted LS with iterated variance")

    g = np.nansum(x * y) / np.nansum(x * x)
    return (float(g), "unweighted LS (no errors provided)")


def _weighted_linreg(x: np.ndarray, y: np.ndarray, weights: np.ndarray):
    weights = np.asarray(weights, dtype=float).ravel()
    weights = np.clip(weights, 1.0e-12, None)
    X = np.vstack([np.ones_like(x), x]).T
    WX = X * weights[:, None]
    XtWX = X.T @ WX
    XtWy = X.T @ (weights * y)
    beta = np.linalg.solve(XtWX, XtWy)
    intercept, slope = beta
    resid = y - (X @ beta)
    dof = builtins.max(len(x) - 2, 1)
    sigma2 = float((weights * resid ** 2).sum() / dof)
    cov = np.linalg.inv(XtWX) * sigma2
    return intercept, slope, cov, resid


def _robust_linfit(
    x: np.ndarray,
    y: np.ndarray,
    yerr: Optional[np.ndarray] = None,
    max_iter: int = 20,
    tol: float = 1.0e-6,
) -> Tuple[float, float, float, float, np.ndarray]:
    """Iteratively reweighted least squares with Huber weights; returns slope/intercept and inlier mask."""
    x = np.asarray(x, dtype=float).ravel()
    y = np.asarray(y, dtype=float).ravel()
    if yerr is not None:
        yerr = np.asarray(yerr, dtype=float).ravel()
        if yerr.size != x.size:
            yerr = None

    n = x.size
    if n < 2:
        return (np.nan, np.nan, np.nan, np.nan, np.zeros(n, dtype=bool))

    base_weights = np.ones_like(x)
    if yerr is not None:
        finite_err = np.isfinite(yerr)
        if np.any(finite_err):
            base_weights[finite_err] = 1.0 / np.clip(yerr[finite_err], 1.0e-8, None) ** 2
            if np.any(~finite_err):
                base_weights[~finite_err] = builtins.max(base_weights[finite_err])
        else:
            yerr = None

    weights = base_weights.copy()
    intercept = np.nan
    slope = np.nan
    cov = None

    for _ in range(max_iter):
        try:
            intercept, slope, cov, resid = _weighted_linreg(x, y, weights)
        except np.linalg.LinAlgError:
            return (np.nan, np.nan, np.nan, np.nan, np.zeros(n, dtype=bool))
        scaled = resid * np.sqrt(weights)
        scale = np.median(np.abs(scaled)) / 0.6745 if scaled.size else 0.0
        if not np.isfinite(scale) or scale <= 1.0e-12:
            scale = np.sqrt(np.average(resid ** 2, weights=weights))
        if not np.isfinite(scale) or scale <= 1.0e-12:
            break
        r = scaled / scale
        c = 1.345
        robust = np.ones_like(r)
        outliers = np.abs(r) > c
        robust[outliers] = c / np.abs(r[outliers])
        new_weights = base_weights * robust
        if np.allclose(new_weights, weights, rtol=tol, atol=tol):
            weights = new_weights
            break
        weights = new_weights

    try:
        intercept, slope, cov, resid = _weighted_linreg(x, y, weights)
    except np.linalg.LinAlgError:
        return (np.nan, np.nan, np.nan, np.nan, np.zeros(n, dtype=bool))

    slope_err = float(np.sqrt(cov[1, 1])) if cov is not None and np.isfinite(cov[1, 1]) else np.nan
    intercept_err = float(np.sqrt(cov[0, 0])) if cov is not None and np.isfinite(cov[0, 0]) else np.nan

    scaled = resid * np.sqrt(weights)
    scale = np.median(np.abs(scaled)) / 0.6745 if scaled.size else 0.0
    if not np.isfinite(scale) or scale <= 1.0e-12:
        scale = np.sqrt(np.average(resid ** 2, weights=weights))
    if not np.isfinite(scale) or scale <= 1.0e-12:
        scale = 1.0
    r = scaled / scale
    c = 1.345
    inlier_mask = np.abs(r) <= c

    return (slope, intercept, slope_err, intercept_err, inlier_mask)


def _plot_flux_comparison(
    ref_flux: np.ndarray,
    test_flux: np.ndarray,
    *,
    xerr: Optional[np.ndarray] = None,
    yerr: Optional[np.ndarray] = None,
    title: str,
    xlabel: str,
    ylabel: str,
    out_path: str,
) -> Dict[str, float]:
    """Scatter plot with robust linear fit, returning fit metadata."""
    plt.figure(figsize=(6.6, 5.5))
    plt.title(title)

    ref_arr = np.asarray(ref_flux, dtype=float).ravel()
    test_arr = np.asarray(test_flux, dtype=float).ravel()
    mask = np.isfinite(ref_arr) & np.isfinite(test_arr)

    if xerr is not None:
        xerr = np.asarray(xerr, dtype=float).ravel()
        if xerr.size == ref_arr.size:
            xerr_valid = xerr[mask]
        else:
            xerr_valid = None
    else:
        xerr_valid = None

    if yerr is not None:
        yerr = np.asarray(yerr, dtype=float).ravel()
        if yerr.size == ref_arr.size:
            yerr_valid = yerr[mask]
        else:
            yerr_valid = None
    else:
        yerr_valid = None

    ref_valid = ref_arr[mask]
    test_valid = test_arr[mask]

    slope = intercept = slope_err = intercept_err = np.nan
    inlier_mask = np.zeros_like(ref_valid, dtype=bool)

    if ref_valid.size:
        # Axis-based sanity check: reject extreme x or y values before linear fitting.
        axis_mask = _axis_inlier_mask(ref_valid) & _axis_inlier_mask(test_valid)
        if np.count_nonzero(axis_mask) >= 2:
            ref_axis = ref_valid[axis_mask]
            test_axis = test_valid[axis_mask]
            yerr_axis = None
            if yerr_valid is not None:
                yerr_axis = yerr_valid[axis_mask]
            # Residual-based screening only runs on axis-filtered points.
            _, _, _, _, residual_mask = _robust_linfit(ref_axis, test_axis, yerr=yerr_axis)
            inlier_mask = np.zeros_like(ref_valid, dtype=bool)
            inlier_mask[axis_mask] = residual_mask
        else:
            axis_mask = np.zeros_like(ref_valid, dtype=bool)
            inlier_mask = np.zeros_like(ref_valid, dtype=bool)

        if np.count_nonzero(inlier_mask) >= 2:
            ref_in = ref_valid[inlier_mask]
            test_in = test_valid[inlier_mask]
            yerr_in = None
            if yerr_valid is not None:
                yerr_in = yerr_valid[inlier_mask]
            # Second pass: fit line using only the inliers so outliers never influence the slope.
            slope, intercept, slope_err, intercept_err, _ = _robust_linfit(
                ref_in, test_in, yerr=yerr_in
            )
        else:
            slope = np.nan
            intercept = np.nan
            slope_err = np.nan
            intercept_err = np.nan
    outlier_mask = ~inlier_mask
    if ref_valid.size:
        vals = np.concatenate([np.abs(ref_valid), np.abs(test_valid)])
        lim = 1.05 * float(np.nanmax(vals)) if vals.size else 1.0
        if not np.isfinite(lim) or lim <= 0.0:
            lim = 1.0
    else:
        lim = 1.0

    if ref_valid.size:
        if np.any(inlier_mask):
            plt.errorbar(
                ref_valid[inlier_mask],
                test_valid[inlier_mask],
                xerr=None if xerr_valid is None else xerr_valid[inlier_mask],
                yerr=None if yerr_valid is None else yerr_valid[inlier_mask],
                fmt=".",
                alpha=0.7,
                capsize=0,
                color="C0",
                label="Data (inliers)",
            )
        if np.any(outlier_mask):
            plt.errorbar(
                ref_valid[outlier_mask],
                test_valid[outlier_mask],
                xerr=None if xerr_valid is None else xerr_valid[outlier_mask],
                yerr=None if yerr_valid is None else yerr_valid[outlier_mask],
                fmt="o",
                alpha=0.8,
                capsize=0,
                color="C3",
                mfc="none",
                label="Data (outliers)",
            )

    plt.plot([0.0, lim], [0.0, lim], "k--", lw=1, label="y = x")

    if np.isfinite(slope) and np.isfinite(intercept):
        x_vals = np.array([0.0, lim])
        y_vals = slope * x_vals + intercept
        if np.any(np.isfinite(y_vals)):
            y_finite = y_vals[np.isfinite(y_vals)]
            if y_finite.size:
                lim = 1.05 * builtins.max(lim, float(np.nanmax(np.abs(y_finite))))
                x_vals = np.array([0.0, lim])
                y_vals = slope * x_vals + intercept
        if np.isfinite(slope_err) and np.isfinite(intercept_err):
            label = (
                f"Fit: y=({slope:.2f}+/-{slope_err:.2f})x + ({intercept:.2g}+/-{intercept_err:.2g})"
            )
        else:
            label = f"Fit: y={slope:.2f}x + {intercept:.2g}"
        plt.plot(x_vals, y_vals, color="C1", lw=1.4, label=label)

    plt.xlabel(xlabel)
    plt.ylabel(ylabel)
    plt.grid(ls=":", alpha=0.5)
    plt.legend(loc="best")
    plt.xlim(0.0, lim)
    plt.ylim(0.0, lim)
    plt.tight_layout()
    _save_fig(out_path)

    return {
        "slope": float(slope),
        "intercept": float(intercept),
        "slope_err": float(slope_err),
        "intercept_err": float(intercept_err),
        "n_inliers": int(np.sum(inlier_mask)),
        "n_outliers": int(np.sum(outlier_mask)),
    }


def compare_fluxes_across_band_and_scans(
    ref_low_xmatch: str,
    ref_high_xmatch: str,
    ref_mfs_xmatch: str,
    *,
    scans_glob: Optional[str] = None,
    docx_name: Optional[str] = None,
) -> Dict[str, object]:
    """
    Compare source fluxes between reference and (low/high/full-band) catalogues and build a report.
    """
    parent = os.path.dirname(os.path.abspath(ref_low_xmatch))
    outdir = _ensure_dir(os.path.join(parent, "crossmatched-fluxes"))
    print(f"\n**** Output directory: {outdir}")

    t_low = Table.read(ref_low_xmatch)
    t_high = Table.read(ref_high_xmatch)
    t_mfs = Table.read(ref_mfs_xmatch)

    # Peak flux comparisons
    R_low_pk = _safe_arr(t_low["Peak_flux_1"])
    T_low_pk = _safe_arr(t_low["Peak_flux_2"])
    Re_low_pk = _col_or_default(t_low, "E_Peak_flux_1", like=R_low_pk)
    Te_low_pk = _col_or_default(t_low, "E_Peak_flux_2", like=T_low_pk)

    R_high_pk = _safe_arr(t_high["Peak_flux_1"])
    T_high_pk = _safe_arr(t_high["Peak_flux_2"])
    Re_high_pk = _col_or_default(t_high, "E_Peak_flux_1", like=R_high_pk)
    Te_high_pk = _col_or_default(t_high, "E_Peak_flux_2", like=T_high_pk)

    R_mfs_pk = _safe_arr(t_mfs["Peak_flux_1"])
    T_mfs_pk = _safe_arr(t_mfs["Peak_flux_2"])
    Re_mfs_pk = _col_or_default(t_mfs, "E_Peak_flux_1", like=R_mfs_pk)
    Te_mfs_pk = _col_or_default(t_mfs, "E_Peak_flux_2", like=T_mfs_pk)

    frac_low_pk = _frac_diff(R_low_pk, T_low_pk)
    frac_high_pk = _frac_diff(R_high_pk, T_high_pk)
    frac_mfs_pk = _frac_diff(R_mfs_pk, T_mfs_pk)

    ratio_low_pk = _ratio(T_low_pk, R_low_pk)
    ratio_high_pk = _ratio(T_high_pk, R_high_pk)
    ratio_mfs_pk = _ratio(T_mfs_pk, R_mfs_pk)

    gain_low_pk = _gain_slope_through_origin(R_low_pk, T_low_pk, Re_low_pk, Te_low_pk)
    gain_high_pk = _gain_slope_through_origin(R_high_pk, T_high_pk, Re_high_pk, Te_high_pk)
    gain_mfs_pk = _gain_slope_through_origin(R_mfs_pk, T_mfs_pk, Re_mfs_pk, Te_mfs_pk)

    # Total (integrated) flux comparisons (optional)
    has_low_total = {"Total_flux_1", "Total_flux_2"}.issubset(t_low.colnames)
    has_high_total = {"Total_flux_1", "Total_flux_2"}.issubset(t_high.colnames)
    has_mfs_total = {"Total_flux_1", "Total_flux_2"}.issubset(t_mfs.colnames)

    if has_low_total:
        R_low_tot = _safe_arr(t_low["Total_flux_1"])
        T_low_tot = _safe_arr(t_low["Total_flux_2"])
        Re_low_tot = _col_or_default(t_low, "E_Total_flux_1", like=R_low_tot)
        Te_low_tot = _col_or_default(t_low, "E_Total_flux_2", like=T_low_tot)
        frac_low_tot = _frac_diff(R_low_tot, T_low_tot)
        ratio_low_tot = _ratio(T_low_tot, R_low_tot)
        gain_low_tot = _gain_slope_through_origin(R_low_tot, T_low_tot, Re_low_tot, Te_low_tot)
    else:
        R_low_tot = T_low_tot = Re_low_tot = Te_low_tot = np.array([])
        frac_low_tot = ratio_low_tot = np.array([])
        gain_low_tot = (np.nan, "Total_flux_* not present in lowband table")

    if has_high_total:
        R_high_tot = _safe_arr(t_high["Total_flux_1"])
        T_high_tot = _safe_arr(t_high["Total_flux_2"])
        Re_high_tot = _col_or_default(t_high, "E_Total_flux_1", like=R_high_tot)
        Te_high_tot = _col_or_default(t_high, "E_Total_flux_2", like=T_high_tot)
        frac_high_tot = _frac_diff(R_high_tot, T_high_tot)
        ratio_high_tot = _ratio(T_high_tot, R_high_tot)
        gain_high_tot = _gain_slope_through_origin(R_high_tot, T_high_tot, Re_high_tot, Te_high_tot)
    else:
        R_high_tot = T_high_tot = Re_high_tot = Te_high_tot = np.array([])
        frac_high_tot = ratio_high_tot = np.array([])
        gain_high_tot = (np.nan, "Total_flux_* not present in highband table")

    if has_mfs_total:
        R_mfs_tot = _safe_arr(t_mfs["Total_flux_1"])
        T_mfs_tot = _safe_arr(t_mfs["Total_flux_2"])
        Re_mfs_tot = _col_or_default(t_mfs, "E_Total_flux_1", like=R_mfs_tot)
        Te_mfs_tot = _col_or_default(t_mfs, "E_Total_flux_2", like=T_mfs_tot)
        frac_mfs_tot = _frac_diff(R_mfs_tot, T_mfs_tot)
        ratio_mfs_tot = _ratio(T_mfs_tot, R_mfs_tot)
        gain_mfs_tot = _gain_slope_through_origin(R_mfs_tot, T_mfs_tot, Re_mfs_tot, Te_mfs_tot)
    else:
        R_mfs_tot = T_mfs_tot = Re_mfs_tot = Te_mfs_tot = np.array([])
        frac_mfs_tot = ratio_mfs_tot = np.array([])
        gain_mfs_tot = (np.nan, "Total_flux_* not present in MFS table")

    # Figure paths
    low_tag = os.path.basename(ref_low_xmatch).replace(".fits", "")
    high_tag = os.path.basename(ref_high_xmatch).replace(".fits", "")
    mfs_tag = os.path.basename(ref_mfs_xmatch).replace(".fits", "")
    if docx_name is None:
        docx_name = f"fluxcmp_{low_tag}__{high_tag}__{mfs_tag}.docx"
    report_docx = os.path.join(outdir, docx_name)

    # Create plots (peak)
    fig_peak_low = os.path.join(outdir, f"ref_vs_low_peak_{low_tag}.png")
    fit_low_pk = _plot_flux_comparison(
        R_low_pk,
        T_low_pk,
        xerr=Re_low_pk,
        yerr=Te_low_pk,
        title="Reference vs Lowband: Peak Flux",
        xlabel="Reference peak flux [Jy/beam]",
        ylabel="Lowband peak flux [Jy/beam]",
        out_path=fig_peak_low,
    )

    fig_peak_high = os.path.join(outdir, f"ref_vs_high_peak_{high_tag}.png")
    fit_high_pk = _plot_flux_comparison(
        R_high_pk,
        T_high_pk,
        xerr=Re_high_pk,
        yerr=Te_high_pk,
        title="Reference vs Highband: Peak Flux",
        xlabel="Reference peak flux [Jy/beam]",
        ylabel="Highband peak flux [Jy/beam]",
        out_path=fig_peak_high,
    )

    fig_peak_mfs = os.path.join(outdir, f"ref_vs_mfs_peak_{mfs_tag}.png")
    fit_mfs_pk = _plot_flux_comparison(
        R_mfs_pk,
        T_mfs_pk,
        xerr=Re_mfs_pk,
        yerr=Te_mfs_pk,
        title="Reference vs Full-band (MFS): Peak Flux",
        xlabel="Reference peak flux [Jy/beam]",
        ylabel="Full-band peak flux [Jy/beam]",
        out_path=fig_peak_mfs,
    )

    fig_peak_hist = os.path.join(outdir, f"fracdiff_hist_peak_{low_tag}_{high_tag}_{mfs_tag}.png")
    plt.figure(figsize=(6.8, 5.5))
    bins = np.linspace(-1.0, 1.0, 60)
    for label, data in [
        ("Lowband", frac_low_pk),
        ("Highband", frac_high_pk),
        ("Full-band MFS", frac_mfs_pk),
    ]:
        finite = data[np.isfinite(data)]
        if finite.size:
            plt.hist(finite, bins=bins, alpha=0.6, label=label, density=False)
    plt.axvline(0, color="k", ls="--", lw=1)
    plt.xlabel("(Test − Ref) / Ref")
    plt.ylabel("Count")
    plt.title("Fractional Peak-Flux Difference")
    plt.legend()
    plt.grid(ls=":", alpha=0.5)
    _save_fig(fig_peak_hist)

    # Total plots
    fig_total_low = fig_total_high = fig_total_mfs = fig_total_hist = None
    fit_total_low = fit_total_high = fit_total_mfs = None

    if has_low_total:
        fig_total_low = os.path.join(outdir, f"ref_vs_low_total_{low_tag}.png")
        fit_total_low = _plot_flux_comparison(
            R_low_tot,
            T_low_tot,
            xerr=Re_low_tot,
            yerr=Te_low_tot,
            title="Reference vs Lowband: Total Flux",
            xlabel="Reference total flux [Jy]",
            ylabel="Lowband total flux [Jy]",
            out_path=fig_total_low,
        )

    if has_high_total:
        fig_total_high = os.path.join(outdir, f"ref_vs_high_total_{high_tag}.png")
        fit_total_high = _plot_flux_comparison(
            R_high_tot,
            T_high_tot,
            xerr=Re_high_tot,
            yerr=Te_high_tot,
            title="Reference vs Highband: Total Flux",
            xlabel="Reference total flux [Jy]",
            ylabel="Highband total flux [Jy]",
            out_path=fig_total_high,
        )

    if has_mfs_total:
        fig_total_mfs = os.path.join(outdir, f"ref_vs_mfs_total_{mfs_tag}.png")
        fit_total_mfs = _plot_flux_comparison(
            R_mfs_tot,
            T_mfs_tot,
            xerr=Re_mfs_tot,
            yerr=Te_mfs_tot,
            title="Reference vs Full-band (MFS): Total Flux",
            xlabel="Reference total flux [Jy]",
            ylabel="Full-band total flux [Jy]",
            out_path=fig_total_mfs,
        )

    if has_low_total or has_high_total or has_mfs_total:
        fig_total_hist = os.path.join(outdir, f"fracdiff_hist_total_{low_tag}_{high_tag}_{mfs_tag}.png")
        plt.figure(figsize=(6.8, 5.5))
        for label, ok, data in [
            ("Lowband", has_low_total, frac_low_tot),
            ("Highband", has_high_total, frac_high_tot),
            ("Full-band MFS", has_mfs_total, frac_mfs_tot),
        ]:
            if ok:
                finite = data[np.isfinite(data)]
                if finite.size:
                    plt.hist(finite, bins=bins, alpha=0.6, label=label, density=False)
        plt.axvline(0, color="k", ls="--", lw=1)
        plt.xlabel("(Test − Ref) / Ref")
        plt.ylabel("Count")
        plt.title("Fractional Total-Flux Difference")
        plt.legend()
        plt.grid(ls=":", alpha=0.5)
        _save_fig(fig_total_hist)

    # Optional per-scan analysis (peak only)
    scan_rows = None
    fig_scan = None
    if scans_glob:
        scan_paths = sorted(glob.glob(scans_glob))
        scan_ids: List[str] = []
        scan_mean: List[float] = []
        scan_std: List[float] = []
        scan_n: List[int] = []
        for path in scan_paths:
            try:
                table = Table.read(path)
                R = _safe_arr(table["Peak_flux_1"])
                T = _safe_arr(table["Peak_flux_2"])
                frac = _frac_diff(R, T)
                scan_ids.append(os.path.basename(path))
                scan_mean.append(float(np.nanmean(frac)))
                scan_std.append(float(np.nanstd(frac)))
                scan_n.append(int(np.isfinite(frac).sum()))
            except Exception:
                scan_ids.append(os.path.basename(path) + " [read_error]")
                scan_mean.append(np.nan)
                scan_std.append(np.nan)
                scan_n.append(0)

        if scan_ids:
            fig_scan = os.path.join(outdir, "per_scan_fracdiff_peak.png")
            plt.figure(figsize=(builtins.max(7.0, 0.25 * len(scan_ids)), 5.2))
            x = np.arange(len(scan_ids))
            plt.errorbar(x, scan_mean, yerr=scan_std, fmt="o", capsize=2)
            plt.axhline(0, color="k", ls="--", lw=1)
            plt.xticks(x, [sid[:30] for sid in scan_ids], rotation=60, ha="right")
            plt.ylabel("(Test − Ref) / Ref")
            plt.title("Per-scan fractional peak-flux difference (mean ± 1σ)")
            plt.grid(ls=":", alpha=0.4)
            _save_fig(fig_scan)

            scan_rows = [["scan_id", "N", "mean_fracdiff", "std_fracdiff"]]
            for sid, n, m, s in zip(scan_ids, scan_n, scan_mean, scan_std):
                scan_rows.append([sid, str(n), f"{m:.6g}", f"{s:.6g}"])

    # Summary rows
    rows_low_pk = _summary_rows("Lowband (Peak)", frac_low_pk, ratio_low_pk, gain_low_pk)
    rows_high_pk = _summary_rows("Highband (Peak)", frac_high_pk, ratio_high_pk, gain_high_pk)
    rows_mfs_pk = _summary_rows("Full-band MFS (Peak)", frac_mfs_pk, ratio_mfs_pk, gain_mfs_pk)

    rows_low_tot = (
        _summary_rows("Lowband (Total)", frac_low_tot, ratio_low_tot, gain_low_tot)
        if has_low_total
        else [["metric", "Lowband (Total): value"], ["note", "Total_flux_* not available"]]
    )
    rows_high_tot = (
        _summary_rows("Highband (Total)", frac_high_tot, ratio_high_tot, gain_high_tot)
        if has_high_total
        else [["metric", "Highband (Total): value"], ["note", "Total_flux_* not available"]]
    )
    rows_mfs_tot = (
        _summary_rows("Full-band MFS (Total)", frac_mfs_tot, ratio_mfs_tot, gain_mfs_tot)
        if has_mfs_total
        else [["metric", "Full-band MFS (Total): value"], ["note", "Total_flux_* not available"]]
    )

    # DOCX report
    doc = Document()
    doc.add_heading("Cross-matched Flux Comparison", level=0)
    doc.add_paragraph(f"Lowband xmatch table:  {os.path.basename(ref_low_xmatch)}")
    doc.add_paragraph(f"Highband xmatch table: {os.path.basename(ref_high_xmatch)}")
    doc.add_paragraph(f"Full-band MFS xmatch table: {os.path.basename(ref_mfs_xmatch)}")
    if scans_glob:
        doc.add_paragraph(f"Scan pattern (peak per-scan): {scans_glob}")
    doc.add_paragraph(
        "Assuming the reference flux is the ground truth, we analyse fractional differences "
        "(test − ref)/ref, flux ratios (test/ref), through-origin gain estimates, and robust "
        "linear fits (with inliers/outliers separated)."
    )

    peak_figures = [
        (
            fig_peak_low,
            "Reference vs Lowband peak flux with robust linear fit (legend lists slope/intercept +/- 1 sigma fit to inliers; dashed line marks y = x).",
        ),
        (
            fig_peak_high,
            "Reference vs Highband peak flux with robust linear fit (legend lists slope/intercept +/- 1 sigma fit to inliers; dashed line marks y = x).",
        ),
        (
            fig_peak_mfs,
            "Reference vs Full-band (MFS) peak flux with robust linear fit (legend lists slope/intercept +/- 1 sigma fit to inliers; dashed line marks y = x).",
        ),
        (
            fig_peak_hist,
            "Histogram of fractional peak-flux differences (test − ref)/ref for Lowband, Highband, and Full-band MFS.",
        ),
    ]
    fig_counter = 1
    for fig_path, desc in peak_figures:
        if fig_path and os.path.exists(fig_path):
            doc.add_picture(fig_path, width=Inches(6.5))
            _docx_add_caption(doc, f"Figure {fig_counter}: {desc}")
            fig_counter += 1

    if fig_scan and os.path.exists(fig_scan):
        doc.add_picture(fig_scan, width=Inches(6.5))
        _docx_add_caption(doc, f"Figure {fig_counter}: Per-scan mean ± 1σ of fractional peak-flux difference.")
        fig_counter += 1

    total_figures = [
        (
            fig_total_low,
            "Figure T1: Reference vs Lowband total flux with robust linear fit (legend lists slope/intercept +/- 1 sigma; dashed line marks y = x).",
        ),
        (
            fig_total_high,
            "Figure T2: Reference vs Highband total flux with robust linear fit (legend lists slope/intercept +/- 1 sigma; dashed line marks y = x).",
        ),
        (
            fig_total_mfs,
            "Figure T3: Reference vs Full-band (MFS) total flux with robust linear fit (legend lists slope/intercept +/- 1 sigma; dashed line marks y = x).",
        ),
        (
            fig_total_hist,
            "Figure T4: Histogram of fractional total-flux differences (test − ref)/ref for Lowband, Highband, and Full-band MFS.",
        ),
    ]
    for fig_path, desc in total_figures:
        if fig_path and os.path.exists(fig_path):
            doc.add_picture(fig_path, width=Inches(6.5))
            _docx_add_caption(doc, desc)

    doc.add_heading("Summary statistics — Peak Flux", level=1)
    _add_table(doc, rows_low_pk, "Table P1: Lowband peak-flux comparison metrics.")
    _add_table(doc, rows_high_pk, "Table P2: Highband peak-flux comparison metrics.")
    _add_table(doc, rows_mfs_pk, "Table P3: Full-band (MFS) peak-flux comparison metrics.")

    doc.add_heading("Summary statistics — Total (Integrated) Flux", level=1)
    _add_table(doc, rows_low_tot, "Table T1: Lowband total-flux comparison metrics.")
    _add_table(doc, rows_high_tot, "Table T2: Highband total-flux comparison metrics.")
    _add_table(doc, rows_mfs_tot, "Table T3: Full-band (MFS) total-flux comparison metrics.")

    if scan_rows:
        doc.add_heading("Per-scan fractional difference (Peak)", level=1)
        _add_table(doc, scan_rows, "Table S1: Per-scan metrics of (test − ref)/ref for peak flux.")

    doc.save(report_docx)

    return {
        "outdir": outdir,
        "report_docx": report_docx,
        "plots_peak": [
            p for p in [fig_peak_low, fig_peak_high, fig_peak_mfs, fig_peak_hist, fig_scan] if p
        ],
        "plots_total": [p for p in [fig_total_low, fig_total_high, fig_total_mfs, fig_total_hist] if p],
        "peak_low_summary": rows_low_pk,
        "peak_high_summary": rows_high_pk,
        "peak_mfs_summary": rows_mfs_pk,
        "total_low_summary": rows_low_tot,
        "total_high_summary": rows_high_tot,
        "total_mfs_summary": rows_mfs_tot,
        "scan_summary": scan_rows,
        "fits": {
            "peak_low": fit_low_pk,
            "peak_high": fit_high_pk,
            "peak_mfs": fit_mfs_pk,
            "total_low": fit_total_low,
            "total_high": fit_total_high,
            "total_mfs": fit_total_mfs,
        },
    }


def _build_arg_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(
        description="Compare fluxes across matched PyBDSF catalogues and generate plots + DOCX summary."
    )
    parser.add_argument(
        "--ref-low-xmatch",
        required=True,
        help=(
            "Low-band crossmatch FITS produced by PyBDSF "
            "(e.g. /Users/samuel/SARAO/imagin_varification/2025/Sky-CrossMatches/"
            "Lband_ref_lowband_Lband_1757723806_lowband.fits)."
        ),
    )
    parser.add_argument(
        "--ref-high-xmatch",
        required=True,
        help=(
            "High-band crossmatch FITS (e.g. "
            "/Users/samuel/SARAO/imagin_varification/2025/Sky-CrossMatches/"
            "Lband_ref_highband_Lband_1757723806_highband.fits)."
        ),
    )
    parser.add_argument(
        "--ref-mfs-xmatch",
        required=True,
        help=(
            "Full-band MFS crossmatch FITS (e.g. "
            "/Users/samuel/SARAO/imagin_varification/2025/Sky-CrossMatches/"
            "refL_mfs_1757723806.SDPAllcorrected.MFS.fits)."
        ),
    )
    parser.add_argument(
        "--scans-glob",
        default=None,
        help=(
            "Optional glob for per-scan crossmatches, e.g. "
            "/Users/samuel/SARAO/imagin_varification/2025/Sky-CrossMatches/"
            "refL_mfs_1757723806_scan*.fits"
        ),
    )
    parser.add_argument(
        "--docx-name",
        default=None,
        help="Optional custom DOCX filename (default auto-generated under crossmatched-fluxes/).",
    )
    return parser


def main(argv: Optional[Sequence[str]] = None) -> int:
    parser = _build_arg_parser()
    args = parser.parse_args(argv)

    result = compare_fluxes_across_band_and_scans(
        args.ref_low_xmatch,
        args.ref_high_xmatch,
        args.ref_mfs_xmatch,
        scans_glob=args.scans_glob,
        docx_name=args.docx_name,
    )

    print("**** DOCX report :", result["report_docx"])
    print("\n**** Peak flux plots  :\n", ",\n ".join(result["plots_peak"]))
    if result["plots_total"]:
        print("\n**** Total/integrated flux plots :\n", ",\n ".join(result["plots_total"]))
    if result["scan_summary"]:
        print("\n**** Scan flux summary table rows:", len(result["scan_summary"]) - 1)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
