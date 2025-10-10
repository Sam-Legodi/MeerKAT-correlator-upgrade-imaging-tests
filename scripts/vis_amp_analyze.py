#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Visibility-amplitude sanity checks for CASA MeasurementSets (.ms/.mms).

Outputs
-------
CSV:
  - perrow_amp_stats.csv          : Per-row metrics (MEAN, RMS, FLAG_FRAC, ANT1, ANT2, SCAN, DDID, SPW, POL)
  - mean_vs_scan.csv              : Median MEAN per SCAN,POL
  - rms_vs_scan.csv               : Median RMS per SCAN,POL
  - flagging_vs_scan.csv          : Mean flag fraction per SCAN,POL
  - flagging_by_channel.csv       : Flag fraction per SPW, CHAN, POL

PNGs:
  - mean_vs_scan_split.png
  - rms_vs_scan_split.png
  - mean_per_antenna_split.png
  - rms_per_antenna_split.png   # NEW
  - mean_vs_baseline_split.png
  - rms_vs_baseline_split.png
  - flagging_vs_scan_split.png
  - flagging_per_antenna_split.png
  - flagging_by_channel.png

If python-docx is available, a concise .docx report is produced.

Notes
-----
- Uses casacore.tables to read MS data. Expects a DATA-like column among ["CORRECTED_DATA", "DATA", "MODEL_DATA"].
- Detrending uses Savitzky–Golay if SciPy is present. Otherwise a robust median removal.
- Split-panel plots separate parallel-hand correlations (XX/YY/LL/RR) from cross-hands (XY/YX/LR/RL).
"""
import os, sys, time, logging, argparse
from pathlib import Path
from typing import Dict, List, Tuple, Optional

import numpy as np
import pandas as pd
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt

# casacore
try:
    from casacore.tables import table as ctable
except Exception as e:
    raise #RuntimeError("casacore.tables is required to read MeasurementSets") from e

# optional Savitzky–Golay
try:
    from scipy.signal import savgol_filter
    HAVE_SAVGOL = True
except Exception:
    HAVE_SAVGOL = False

# optional docx
try:
    from docx import Document
    from docx.shared import Inches
    HAVE_DOCX = True
except Exception:
    HAVE_DOCX = False

logging.basicConfig(level=logging.INFO, format="%(asctime)s [%(levelname)s] %(message)s")
logger = logging.getLogger(__name__)

# ------------------------ Environment helpers ------------------------

def ensure_screen_or_tmux() -> None:
    if os.environ.get("ALLOW_NO_TMUX", "").lower() in ("1", "y", "yes", "true"):
        return
    if ("STY" in os.environ) or ("TMUX" in os.environ):
        return
    ans = os.environ.get("NONINTERACTIVE_ACK", "").lower()
    if ans in ("y", "yes", "1", "true"):
        return
    try:
        reply = input(
            "[WARN] Not running in screen/tmux. Continue anyway? [y/N] "
            "(set ALLOW_NO_TMUX=1 to suppress this check): "
        ).strip().lower()
    except EOFError:
        reply = "n"
    if reply not in ("y", "yes"):
        sys.exit(1)

# ------------------------ MS helpers ------------------------

def _pick_data_column(ms) -> str:
    cols = ms.colnames()
    for c in ("CORRECTED_DATA", "DATA", "MODEL_DATA"):
        if c in cols:
            return c
    raise RuntimeError("No DATA-like column found among CORRECTED_DATA/DATA/MODEL_DATA")

def _get_field_id(ms_path: str, field_sel: Optional[str]) -> Tuple[Optional[int], Optional[str]]:
    """Return (FIELD_ID, FIELD_NAME). If field_sel is None, returns (None, None) for all fields."""
    if field_sel is None:
        return None, None
    with ctable(f"{ms_path}::FIELD", readonly=True) as ft:
        names = ft.getcol("NAME")
    # exact name match
    for i, n in enumerate(names):
        if n == str(field_sel):
            return i, n
    # integer index
    try:
        fid = int(field_sel)
    except Exception:
        raise RuntimeError(f'Field "{field_sel}" not found. Available: {list(names)}')
    if 0 <= fid < len(names):
        return fid, names[fid]
    raise RuntimeError(f'Field id {fid} out of range 0..{len(names)-1}')

def _ddid_to_spw_map(ms_path: str) -> np.ndarray:
    with ctable(f"{ms_path}::DATA_DESCRIPTION", readonly=True) as dd:
        return dd.getcol("SPECTRAL_WINDOW_ID")

def _spw_chan_freqs(ms_path: str) -> Tuple[np.ndarray, np.ndarray]:
    with ctable(f"{ms_path}::SPECTRAL_WINDOW", readonly=True) as spw:
        return spw.getcol("CHAN_FREQ"), spw.getcol("NUM_CHAN")

def _pol_labels(ms_path: str) -> List[str]:
    with ctable(f"{ms_path}::POLARIZATION", readonly=True) as pt:
        corr_types = pt.getcol("CORR_TYPE")  # list of arrays
    # Use first row
    if len(corr_types) == 0:
        return []
    lookup = {
        5:"RR", 6:"RL", 7:"LR", 8:"LL",
        9:"XX", 10:"XY", 11:"YX", 12:"YY",
    }
    return [lookup.get(int(x), str(int(x))) for x in corr_types[0]]

def _is_auto(ant1: int, ant2: int) -> bool:
    return int(ant1) == int(ant2)

def _ensure_odd(n: int) -> int:
    return n if n % 2 == 1 else n + 1

def _detrend_amp(chan_amp: np.ndarray, window: int, order: int) -> np.ndarray:
    """Return residual = amp - trend."""
    if HAVE_SAVGOL and window >= order + 2 and window >= 5:
        w = _ensure_odd(max(5, min(window, len(chan_amp) - (len(chan_amp)%2==0))))
        if w < order + 2:
            w = order + 3
        w = _ensure_odd(w)
        try:
            trend = savgol_filter(chan_amp, window_length=w, polyorder=min(order, w-2), mode="interp")
            return chan_amp - trend
        except Exception:
            pass
    med = np.nanmedian(chan_amp)
    return chan_amp - med

# ------------------------ Core analysis ------------------------

def analyze_single(ms_path: str,
                   field_sel: Optional[str]=None,
                   outdir: Optional[str]=None,
                   window: int=51,
                   order: int=3,
                   chan_min: Optional[int]=None,
                   chan_max: Optional[int]=None) -> Dict[str, Path]:
    """
    Compute per-row mean amplitude and RMS of detrended spectrum.
    Also aggregate by scan, antenna, and baseline. Write CSVs and plots.

    Parameters
    ----------
    ms_path : str
    field_sel : str|None
    outdir   : str|None
    window   : int  Savitzky–Golay window (odd). Fallback to median when SciPy missing.
    order    : int  Polynomial order for Savitzky–Golay.
    chan_min/max : Optional[int] channel slice within each SPW.

    Returns
    -------
    dict with paths and the created outdir.
    """
    ms_path = str(ms_path)
    ms_path_p = Path(ms_path)
    if outdir is None:
        base = ms_path_p.with_suffix("").name
        outdir = ms_path_p.parent / f"{base}_output"
    outdir = Path(outdir)
    if outdir.exists() and any(outdir.iterdir()):
        # disambiguate
        ts = time.strftime("%Y-%m-%d_%H.%M")
        outdir = Path(f"{outdir}_{ts}")
    outdir.mkdir(parents=True, exist_ok=True)
    logger.info(f"Output: {outdir}")

    # polarization labels
    pol_labels = _pol_labels(ms_path)
    npol = len(pol_labels)
    parallel = [p for p in pol_labels if p in ("XX","YY","LL","RR")]
    cross    = [p for p in pol_labels if p in ("XY","YX","LR","RL")]

    # row selection
    fid, field_name = _get_field_id(ms_path, field_sel)
    with ctable(ms_path, readonly=True) as ms:
        msq = ms if fid is None else ms.query(f"FIELD_ID=={fid}")
        logger.info(f"Field restriction: {field_name if field_name else 'ALL'}")
        data_col = _pick_data_column(msq)
        have_scan = "SCAN_NUMBER" in msq.colnames()
        have_flag = "FLAG" in msq.colnames()

        # spw info
        chan_freqs, num_chan = _spw_chan_freqs(ms_path)   # list-like
        dd2spw = _ddid_to_spw_map(ms_path)

        # bounds per spw
        spw_bounds: Dict[int, Tuple[int,int]] = {}
        for spw_id, n in enumerate(num_chan):
            lo = 0 if chan_min is None else max(0, int(chan_min))
            hi = n if chan_max is None else min(n, int(chan_max))
            if lo >= hi:
                lo, hi = 0, n
            spw_bounds[spw_id] = (lo, hi)

        # accumulators
        rows = []
        # flag stats by channel and pol class
        flag_counts = {cls: {spw:int(n)*np.zeros((int(num_chan[spw]), npol), dtype=np.int64)
                             for spw in range(len(num_chan))}
                       for cls in ("auto","cross")}
        tot_counts  = {cls: {spw:int(n)*np.zeros((int(num_chan[spw]), npol), dtype=np.int64)
                             for spw in range(len(num_chan))}
                       for cls in ("auto","cross")}

        nrows = msq.nrows()
        step = 2048
        for start in range(0, nrows, step):
            stop = min(nrows, start + step)
            sl = slice(start, stop)
            d    = msq.getcol(data_col, startrow=start, nrow=stop-start)          # (nrow, nchan, npol) complex
            ant1 = msq.getcol("ANTENNA1", startrow=start, nrow=stop-start)
            ant2 = msq.getcol("ANTENNA2", startrow=start, nrow=stop-start)
            ddid = msq.getcol("DATA_DESC_ID", startrow=start, nrow=stop-start)
            scan = msq.getcol("SCAN_NUMBER", startrow=start, nrow=stop-start) if have_scan else np.zeros(len(ant1), dtype=int)
            flg  = msq.getcol("FLAG", startrow=start, nrow=stop-start) if have_flag else np.zeros_like(d, dtype=bool)

            # loop rows in this chunk
            for i in range(d.shape[0]):
                spw = int(dd2spw[int(ddid[i])])
                lo, hi = spw_bounds[spw]
                spec = d[i, lo:hi, :]
                if spec.size == 0:
                    continue
                amps = np.abs(spec)  # (nchan', npol)
                if have_flag:
                    f = flg[i, lo:hi, :]
                else:
                    f = np.zeros_like(amps, dtype=bool)
                # flag fraction per pol
                with np.errstate(invalid="ignore", divide="ignore"):
                    flag_frac_pol = np.sum(f, axis=0) / np.maximum(1, amps.shape[0])
                # detrend and RMS per pol
                rms = np.full((amps.shape[1],), np.nan, dtype=float)
                mean_amp = np.nanmean(amps, axis=0)
                for p in range(amps.shape[1]):
                    a = amps[:, p]
                    residual = _detrend_amp(a, window=window, order=order)
                    rms[p] = np.nanstd(residual)

                cls = "auto" if _is_auto(ant1[i], ant2[i]) else "cross"
                # update per-channel flagging counters
                for p in range(npol):
                    tot_counts[cls][spw][lo:hi, p] += 1
                    if have_flag:
                        flag_counts[cls][spw][lo:hi, p] += f[:, p].astype(np.int64)

                # emit one row per pol
                for p in range(npol):
                    rows.append(dict(
                        ANT1=int(ant1[i]), ANT2=int(ant2[i]), BASE=f"{int(ant1[i])}-{int(ant2[i])}",
                        SCAN=int(scan[i]), DDID=int(ddid[i]), SPW=spw, POL=pol_labels[p],
                        MEAN=float(mean_amp[p]), RMS=float(rms[p]),
                        FLAG_FRAC=float(flag_frac_pol[p])
                    ))

    df = pd.DataFrame(rows)
    csv_main = outdir / "perrow_amp_stats.csv"
    df.to_csv(csv_main, index=False)

    # aggregate by scan
    g_scan_mean = df.groupby(["SCAN","POL"])["MEAN"].median().reset_index()
    g_scan_rms  = df.groupby(["SCAN","POL"])["RMS"].median().reset_index()
    g_scan_flag = df.groupby(["SCAN","POL"])["FLAG_FRAC"].mean().reset_index()
    g_scan_mean.to_csv(outdir/"mean_vs_scan.csv", index=False)
    g_scan_rms.to_csv(outdir/"rms_vs_scan.csv", index=False)
    g_scan_flag.to_csv(outdir/"flagging_vs_scan.csv", index=False)

    # aggregate by antenna
    # a per-antenna metric is the median of all baselines touching the antenna
    # Use BOTH ant1 and ant2 contributions
    dfA1 = df.rename(columns={"ANT1":"ANT"}).drop(columns=["ANT2"])
    dfA2 = df.rename(columns={"ANT2":"ANT"}).drop(columns=["ANT1"])
    dfAnt = pd.concat([dfA1[["ANT","POL","MEAN","RMS","FLAG_FRAC"]],
                       dfA2[["ANT","POL","MEAN","RMS","FLAG_FRAC"]]], ignore_index=True)
    g_ant_mean = dfAnt.groupby(["ANT","POL"])["MEAN"].median().reset_index()
    g_ant_rms  = dfAnt.groupby(["ANT","POL"])["RMS"].median().reset_index()
    g_ant_flag = dfAnt.groupby(["ANT","POL"])["FLAG_FRAC"].mean().reset_index()

    # aggregate by baseline
    g_bl_mean = df.groupby(["BASE","POL"])["MEAN"].median().reset_index()
    g_bl_rms  = df.groupby(["BASE","POL"])["RMS"].median().reset_index()

    # flag by channel output
    # need num_chan and counters; rebuild quickly from saved local scope
    # (These variables are still in scope: num_chan, flag_counts, tot_counts, pol_labels)
    rows_chan = []
    for cls in ("auto","cross"):
        for spw in range(len(num_chan)):
            for p in range(npol):
                tot = tot_counts[cls][spw][:, p]
                flg = flag_counts[cls][spw][:, p]
                for ch in range(len(tot)):
                    if tot[ch] == 0:
                        continue
                    rows_chan.append(dict(CLASS=cls, SPW=spw, CHAN=ch, POL=pol_labels[p],
                                          TOTAL=int(tot[ch]), FLAGGED=int(flg[ch]),
                                          FLAG_FRAC=float(flg[ch])/float(tot[ch])))
    pd.DataFrame(rows_chan).to_csv(outdir/"flagging_by_channel.csv", index=False)

    # ------------------------ Plots ------------------------

    def _two_panel(title: str, left_lines, right_lines, xlabel: str, ylabel: str, outname: str) -> Path:
        fig, (axL, axR) = plt.subplots(1, 2, figsize=(14, 5), constrained_layout=True)
        for lab, x, y in left_lines:
            axL.plot(x, y, marker="o", ms=3, lw=1, label=lab)
        for lab, x, y in right_lines:
            axR.plot(x, y, marker="o", ms=3, lw=1, label=lab)
        axL.set_title(f"{title} (parallel)")
        axR.set_title(f"{title} (cross)")
        axL.set_xlabel(xlabel); axR.set_xlabel(xlabel)
        axL.set_ylabel(ylabel); axR.set_ylabel(ylabel)
        axL.grid(ls=":"); axR.grid(ls=":")
        if left_lines: axL.legend()
        if right_lines: axR.legend()
        pth = outdir / outname
        fig.savefig(pth, dpi=150)
        plt.close(fig)
        return pth

    # mean vs scan
    L = [(p,
          g_scan_mean[g_scan_mean.POL==p].sort_values("SCAN")["SCAN"].values,
          g_scan_mean[g_scan_mean.POL==p].sort_values("SCAN")["MEAN"].values)
         for p in parallel if p in g_scan_mean.POL.unique()]
    R = [(p,
          g_scan_mean[g_scan_mean.POL==p].sort_values("SCAN")["SCAN"].values,
          g_scan_mean[g_scan_mean.POL==p].sort_values("SCAN")["MEAN"].values)
         for p in cross if p in g_scan_mean.POL.unique()]
    fig_mean_vs_scan = _two_panel("Mean amplitude vs scan", L, R, "Scan", "Median mean |V|", "mean_vs_scan_split.png")

    # rms vs scan
    L = [(p,
          g_scan_rms[g_scan_rms.POL==p].sort_values("SCAN")["SCAN"].values,
          g_scan_rms[g_scan_rms.POL==p].sort_values("SCAN")["RMS"].values)
         for p in parallel if p in g_scan_rms.POL.unique()]
    R = [(p,
          g_scan_rms[g_scan_rms.POL==p].sort_values("SCAN")["SCAN"].values,
          g_scan_rms[g_scan_rms.POL==p].sort_values("SCAN")["RMS"].values)
         for p in cross if p in g_scan_rms.POL.unique()]
    fig_rms_vs_scan = _two_panel("RMS residual vs scan", L, R, "Scan", "Median RMS", "rms_vs_scan_split.png")

    # per-antenna mean bars
    ants = sorted(g_ant_mean.ANT.unique())
    xidx = np.arange(len(ants))
    fig, (axL, axR) = plt.subplots(1, 2, figsize=(14,5), constrained_layout=True)
    par_pols = [p for p in parallel if p in g_ant_mean.POL.unique()]
    widthL = 0.8 / max(1, len(par_pols))
    for j, p in enumerate(par_pols):
        gg = g_ant_mean[g_ant_mean.POL==p]
        y = [gg[gg.ANT==a].MEAN.median() if (gg.ANT==a).any() else np.nan for a in ants]
        axL.bar(xidx + j*widthL, y, width=widthL, label=p)
    axL.set_xticks(xidx + widthL*max(0.5, len(par_pols)/2.0))
    axL.set_xticklabels([str(a) for a in ants], rotation=90)
    axL.set_title("Per-antenna mean (parallel)")
    axL.set_xlabel("Antenna ID"); axL.set_ylabel("Median mean |V|"); axL.grid(ls=":"); axL.legend()

    cro_pols = [p for p in cross if p in g_ant_mean.POL.unique()]
    widthR = 0.8 / max(1, len(cro_pols))
    for j, p in enumerate(cro_pols):
        gg = g_ant_mean[g_ant_mean.POL==p]
        y = [gg[gg.ANT==a].MEAN.median() if (gg.ANT==a).any() else np.nan for a in ants]
        axR.bar(xidx + j*widthR, y, width=widthR, label=p)
    axR.set_xticks(xidx + widthR*max(0.5, len(cro_pols)/2.0))
    axR.set_xticklabels([str(a) for a in ants], rotation=90)
    axR.set_title("Per-antenna mean (cross)")
    axR.set_xlabel("Antenna ID"); axR.set_ylabel("Median mean |V|"); axR.grid(ls=":"); axR.legend()
    fig_mean_per_ant = outdir/"mean_per_antenna_split.png"
    fig.savefig(fig_mean_per_ant, dpi=150); plt.close(fig)

    # per-antenna RMS bars  (NEW)
    ants = sorted(g_ant_rms.ANT.unique())
    xidx = np.arange(len(ants))
    fig, (axL, axR) = plt.subplots(1, 2, figsize=(14,5), constrained_layout=True)
    par_pols = [p for p in parallel if p in g_ant_rms.POL.unique()]
    widthL = 0.8 / max(1, len(par_pols))
    for j, p in enumerate(par_pols):
        gg = g_ant_rms[g_ant_rms.POL==p]
        y = [gg[gg.ANT==a].RMS.median() if (gg.ANT==a).any() else np.nan for a in ants]
        axL.bar(xidx + j*widthL, y, width=widthL, label=p)
    axL.set_xticks(xidx + widthL*max(0.5, len(par_pols)/2.0))
    axL.set_xticklabels([str(a) for a in ants], rotation=90)
    axL.set_title("Per-antenna RMS (parallel)")
    axL.set_xlabel("Antenna ID"); axL.set_ylabel("Median RMS"); axL.grid(ls=":"); axL.legend()

    cro_pols = [p for p in cross if p in g_ant_rms.POL.unique()]
    widthR = 0.8 / max(1, len(cro_pols))
    for j, p in enumerate(cro_pols):
        gg = g_ant_rms[g_ant_rms.POL==p]
        y = [gg[gg.ANT==a].RMS.median() if (gg.ANT==a).any() else np.nan for a in ants]
        axR.bar(xidx + j*widthR, y, width=widthR, label=p)
    axR.set_xticks(xidx + widthR*max(0.5, len(cro_pols)/2.0))
    axR.set_xticklabels([str(a) for a in ants], rotation=90)
    axR.set_title("Per-antenna RMS (cross)")
    axR.set_xlabel("Antenna ID"); axR.set_ylabel("Median RMS"); axR.grid(ls=":"); axR.legend()
    fig_rms_per_ant = outdir/"rms_per_antenna_split.png"
    fig.savefig(fig_rms_per_ant, dpi=150); plt.close(fig)

    # baseline bars
    bls = sorted(g_bl_mean.BASE.unique())
    xidx = np.arange(len(bls))
    def _bar_by(df_in, value_col: str, outname: str, title: str) -> Path:
        fig, (axL, axR) = plt.subplots(1,2, figsize=(16,6), constrained_layout=True)
        par = [p for p in parallel if p in df_in.POL.unique()]
        widthL = 0.8 / max(1, len(par))
        for j,p in enumerate(par):
            gg = df_in[df_in.POL==p]
            y = [gg[gg.BASE==b][value_col].median() if (gg.BASE==b).any() else np.nan for b in bls]
            axL.bar(xidx + j*widthL, y, width=widthL, label=p)
        axL.set_xticks(xidx + widthL*max(0.5, len(par)/2.0))
        axL.set_xticklabels([str(b) for b in bls], rotation=90)
        axL.set_title(f"{title} (parallel)"); axL.set_xlabel("Baseline"); axL.set_ylabel(value_col); axL.grid(ls=":"); axL.legend()

        cro = [p for p in cross if p in df_in.POL.unique()]
        widthR = 0.8 / max(1, len(cro))
        for j,p in enumerate(cro):
            gg = df_in[df_in.POL==p]
            y = [gg[gg.BASE==b][value_col].median() if (gg.BASE==b).any() else np.nan for b in bls]
            axR.bar(xidx + j*widthR, y, width=widthR, label=p)
        axR.set_xticks(xidx + widthR*max(0.5, len(cro)/2.0))
        axR.set_xticklabels([str(b) for b in bls], rotation=90)
        axR.set_title(f"{title} (cross)"); axR.set_xlabel("Baseline"); axR.set_ylabel(value_col); axR.grid(ls=":"); axR.legend()
        pth = outdir/outname
        fig.savefig(pth, dpi=150); plt.close(fig); return pth

    fig_mean_vs_bl = _bar_by(g_bl_mean, "MEAN", "mean_vs_baseline_split.png", "Per-baseline mean")
    fig_rms_vs_bl  = _bar_by(g_bl_rms,  "RMS",  "rms_vs_baseline_split.png",  "Per-baseline RMS")

    # flagging plots
    # vs scan
    L = [(p,
          g_scan_flag[g_scan_flag.POL==p].sort_values("SCAN")["SCAN"].values,
          g_scan_flag[g_scan_flag.POL==p].sort_values("SCAN")["FLAG_FRAC"].values)
         for p in parallel if p in g_scan_flag.POL.unique()]
    R = [(p,
          g_scan_flag[g_scan_flag.POL==p].sort_values("SCAN")["SCAN"].values,
          g_scan_flag[g_scan_flag.POL==p].sort_values("SCAN")["FLAG_FRAC"].values)
         for p in cross if p in g_scan_flag.POL.unique()]
    fig_flag_vs_scan = _two_panel("Flagging vs scan", L, R, "Scan", "Flagging fraction", "flagging_vs_scan_split.png")

    # per-antenna flag bars
    ants = sorted(g_ant_flag.ANT.unique())
    xidx = np.arange(len(ants))
    fig, (axL, axR) = plt.subplots(1, 2, figsize=(14,5), constrained_layout=True)
    par = [p for p in parallel if p in g_ant_flag.POL.unique()]
    widthL = 0.8 / max(1, len(par))
    for j,p in enumerate(par):
        gg = g_ant_flag[g_ant_flag.POL==p]
        y = [gg[gg.ANT==a].FLAG_FRAC.mean() if (gg.ANT==a).any() else 0.0 for a in ants]
        axL.bar(xidx + j*widthL, y, width=widthL, label=p)
    axL.set_xticks(xidx + widthL*max(0.5, len(par)/2.0))
    axL.set_xticklabels([str(a) for a in ants], rotation=90)
    axL.set_title("Per-antenna flagging (parallel)")
    axL.set_xlabel("Antenna ID"); axL.set_ylabel("Flagging fraction"); axL.grid(ls=":"); axL.legend()

    cro = [p for p in cross if p in g_ant_flag.POL.unique()]
    widthR = 0.8 / max(1, len(cro))
    for j,p in enumerate(cro):
        gg = g_ant_flag[g_ant_flag.POL==p]
        y = [gg[gg.ANT==a].FLAG_FRAC.mean() if (gg.ANT==a).any() else 0.0 for a in ants]
        axR.bar(xidx + j*widthR, y, width=widthR, label=p)
    axR.set_xticks(xidx + widthR*max(0.5, len(cro)/2.0))
    axR.set_xticklabels([str(a) for a in ants], rotation=90)
    axR.set_title("Per-antenna flagging (cross)")
    axR.set_xlabel("Antenna ID"); axR.set_ylabel("Flagging fraction"); axR.grid(ls=":"); axR.legend()
    fig_flag_per_ant = outdir/"flagging_per_antenna_split.png"
    fig.savefig(fig_flag_per_ant, dpi=150); plt.close(fig)

    # flagging by channel heatmap-like line per pol class
    flag_by_chan = pd.read_csv(outdir/"flagging_by_channel.csv")
    def _plot_flag_by_channel(df_in: pd.DataFrame) -> Path:
        fig, (axL, axR) = plt.subplots(1,2, figsize=(14,5), constrained_layout=True)
        for ax, cls in ((axL, "auto"), (axR, "cross")):
            sub = df_in[df_in.CLASS==cls]
            if sub.empty:
                ax.set_title(f"{cls} (no data)"); continue
            # average over SPW for each CHAN,POL
            grp = sub.groupby(["CHAN","POL"])["FLAG_FRAC"].mean().reset_index()
            for p in grp.POL.unique():
                gg = grp[grp.POL==p].sort_values("CHAN")
                ax.plot(gg["CHAN"].values, gg["FLAG_FRAC"].values, lw=1, label=p)
            ax.set_title(f"Flagging by channel ({cls})")
            ax.set_xlabel("Channel"); ax.set_ylabel("Flagging fraction"); ax.grid(ls=":"); ax.legend()
        pth = outdir/"flagging_by_channel.png"
        fig.savefig(pth, dpi=150); plt.close(fig); return pth
    fig_flag_vs_chan = _plot_flag_by_channel(flag_by_chan)

    # optional docx
    if HAVE_DOCX:
        try:
            doc = Document()
            doc.add_heading("Visibility amplitude analysis", level=1)
            doc.add_paragraph(f"MS: {ms_path}")
            if field_name:
                doc.add_paragraph(f"Field: {field_name}")
            doc.add_paragraph(f"Rows processed: {len(df)} (per-pol)")
            # quick KPI table
            t = doc.add_table(rows=1, cols=4)
            hdr = t.rows[0].cells
            hdr[0].text = "Metric"; hdr[1].text = "Group"; hdr[2].text = "Value"; hdr[3].text = "Notes"
            try:
                row = t.add_row().cells
                row[0].text = "Overall flagging"
                row[1].text = "mean over rows"
                row[2].text = f"{df.FLAG_FRAC.mean():.4f}"
                row[3].text = "Lower is better"
            except Exception:
                pass
            doc.add_heading("Figures", level=2)
            for figp in [fig_mean_vs_scan, fig_rms_vs_scan, fig_mean_per_ant, fig_rms_per_ant,
                         fig_mean_vs_bl, fig_rms_vs_bl, fig_flag_vs_scan, fig_flag_per_ant, fig_flag_vs_chan]:
                doc.add_paragraph(Path(figp).name)
                doc.add_picture(str(figp), width=Inches(6.5))
            rep = outdir/"vis_amp_summary.docx"
            doc.save(str(rep))
        except Exception as e:
            logger.warning(f"Could not create Word report: {e}")
    else:
        logger.info("python-docx not available. Skipping .docx report.")

    return {
        "outdir": outdir,
        "csv_main": csv_main,
        "fig_mean_vs_scan": fig_mean_vs_scan,
        "fig_rms_vs_scan": fig_rms_vs_scan,
        "fig_mean_per_ant": fig_mean_per_ant,
        "fig_rms_per_ant": fig_rms_per_ant,
        "fig_mean_vs_bl": fig_mean_vs_bl,
        "fig_rms_vs_bl": fig_rms_vs_bl,
        "fig_flag_vs_scan": fig_flag_vs_scan,
        "fig_flag_per_ant": fig_flag_per_ant,
    }

# ------------------------ Comparison helper (optional) ------------------------

def compare_results(res_a: Dict[str, Path], res_b: Dict[str, Path], outdir: Path) -> None:
    """If both have perrow_amp_stats.csv, produce simple side-by-side scan medians."""
    try:
        a = pd.read_csv(res_a["outdir"]/ "perrow_amp_stats.csv")
        b = pd.read_csv(res_b["outdir"]/ "perrow_amp_stats.csv")
    except Exception as e:
        logger.warning(f"Comparison skipped: {e}")
        return
    for metric in ("MEAN","RMS"):
        ga = a.groupby(["SCAN","POL"])[metric].median().reset_index().rename(columns={metric:f"{metric}_A"})
        gb = b.groupby(["SCAN","POL"])[metric].median().reset_index().rename(columns={metric:f"{metric}_B"})
        m = pd.merge(ga, gb, on=["SCAN","POL"], how="outer").sort_values(["POL","SCAN"])
        m.to_csv(outdir/f"compare_{metric.lower()}_vs_scan.csv", index=False)

# ------------------------ CLI ------------------------

def main():
    ensure_screen_or_tmux()
    ap = argparse.ArgumentParser(description="Analyze visibility amplitude oscillations in a MeasurementSet.")
    ap.add_argument("--ms", required=True, help="Path to .ms or .mms")
    ap.add_argument("--ms-ref", default=None, help="Optional reference .ms/.mms to compare")
    ap.add_argument("--field", default=None, help="FIELD name or integer FIELD_ID")
    ap.add_argument("--outdir", default=None, help="Output directory; default <ms>_output next to the MS")
    ap.add_argument("--window", type=int, default=51, help="Savitzky–Golay(SG) window length (odd). Default 51")
    ap.add_argument("--order", type=int, default=3, help="Polynomial order for SG detrend. Default 3")
    ap.add_argument("--chan-min", type=int, default=None, help="Channel lower bound (inclusive)")
    ap.add_argument("--chan-max", type=int, default=None, help="Channel upper bound (exclusive)")
    args = ap.parse_args()

    try:
        res_test = analyze_single(
            ms_path=args.ms,
            field_sel=args.field,
            outdir=args.outdir,
            window=args.window,
            order=args.order,
            chan_min=args.chan_min,
            chan_max=args.chan_max,
        )
        if args.ms_ref:
            res_ref = analyze_single(
                ms_path=args.ms_ref,
                field_sel=args.field,
                outdir=None,
                window=args.window,
                order=args.order,
                chan_min=args.chan_min,
                chan_max=args.chan_max,
            )
            compare_results(res_test, res_ref, outdir=res_test["outdir"])
        logger.info("Done.")
    except Exception as e:
        logger.error(f"Error: {e}")
        sys.exit(2)

if __name__ == "__main__":
    main()
