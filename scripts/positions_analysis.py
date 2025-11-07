#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Cross-matched positional analysis with optional per-scan overlays.

WHAT THIS DOES
--------------
1) Reads a cross-match FITS table (reference x other).
2) Builds ΔRA/ΔDec (east+, north+) using spherical geometry.
3) Produces 5 figures (pos_varCMC1xCMC2) and a DOCX report.
4) Optionally overlays *per-scan* cross-match tables on the same axes.
5) For any plot with >1 dataset on the same axes, draws an **enclosing ellipse**
   for each dataset:
     - Center = (mean x, mean y)
     - Width  = 2 * max |x - mean_x|
     - Height = 2 * max |y - mean_y|
     - Color  = matches the dataset’s line/marker color
     - A matching colored 'x' marks the center.
6) Legends are placed **outside** the axes (to the right), so markers aren’t obscured.
7) DOCX report includes **physical interpretation** text for each plot.

HOW TO RUN (CLI)
----------------
python 2025_positions_analysis.py \
  --xmatch-table /path/to/refXother.fits \
  --ref-fits /path/to/reference_image.fits \
  --other-fits /path/to/other_image.fits \
  --per-scan-glob "/path/to/per_scan/refXscan*.fits"

Outputs go to: <dir_of_other_fits_or_table>/crossmatched-positions

HOW TO RUN (JUPYTER)
--------------------
import 2025_positions_analysis as posan   # if file name starts with digits, rename to positions_analysis_2025.py
res = posan.analyze_ref_vs_other_with_optional_scans(
    xmatch_table_path="/path/to/refXother.fits",
    ref_fits_path="/path/to/reference_image.fits",
    other_fits_path="/path/to/other_image.fits",
    per_scan_glob="/path/to/per_scan/refXscan*.fits",  # or None
)
print(res["outdir"], res["report_docx"])
"""

import os, sys, time, argparse, csv, glob, builtins
from astropy.table import Table, Column
from astropy.io import fits
from astropy.wcs import WCS as w
from matplotlib import pyplot as plt
from matplotlib.pyplot import *
from astropy.wcs import WCS
import pandas as pd
from math import *
from numpy import *
import itertools
from astropy.coordinates import SkyCoord
import astropy.units as u
import numpy as np
from numpy import pi
from docx import Document
from docx.shared import Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from matplotlib.patches import Ellipse  # NEW: for enclosing ellipses


def safe_div(a, b, default=np.nan):
    try:
        return a / b
    except Exception:
        return default

# === HARMONIC QUADRUPOLE CHECK (paste near helpers) ===
def _harmonic_quad_summary(phi_deg, y, label=""):
    """
    Fit y(φ) with dipole (1φ) and dipole+quadrupole (1φ+2φ) harmonics.
    Returns dict with A2 (amplitude of 2φ), psi2 (phase of 2φ/2 in degrees),
    and ΔR2 = R2(full) - R2(dipole) as the variance uniquely explained by 2φ.

    y and phi_deg can contain NaNs; pairs with NaN are dropped.
    """
    phi = np.asarray(phi_deg, dtype=float)
    y   = np.asarray(y,       dtype=float)
    m   = np.isfinite(phi) & np.isfinite(y)
    phi = phi[m]; y = y[m]
    if y.size < 5:
        return dict(label=label, A2=np.nan, psi2=np.nan, dR2=np.nan, R2_full=np.nan)

    # Radians
    ph = np.deg2rad(phi)

    # Design matrices
    # Dipole only: [1, cosφ, sinφ]
    X1 = np.column_stack([np.ones_like(ph), np.cos(ph), np.sin(ph)])
    # Dipole + Quadrupole: [1, cosφ, sinφ, cos2φ, sin2φ]
    X2 = np.column_stack([X1, np.cos(2*ph), np.sin(2*ph)])

    def _fit(X, y):
        beta, *_ = np.linalg.lstsq(X, y, rcond=None)
        yhat = X @ beta
        resid = y - yhat
        rss = float(np.sum(resid**2))
        tss = float(np.sum((y - np.nanmean(y))**2))
        R2  = np.nan if tss == 0 else 1.0 - rss/tss
        return beta, yhat, rss, tss, R2

    b1, y1, rss1, tss, R2_dip = _fit(X1, y)
    b2, y2, rss2, _,   R2_full = _fit(X2, y)

    # Quadrupole amplitude and phase
    a2 = float(b2[-2])  # cos 2φ
    b2s = float(b2[-1]) # sin 2φ
    A2 = np.hypot(a2, b2s)                         # amplitude of 2φ term
    delta = np.rad2deg(np.arctan2(b2s, a2))        # phase of 2φ (in deg)
    psi2 = 0.5 * delta                             # maxima occur near φ ≈ psi2

    dR2 = (R2_full - R2_dip) if (np.isfinite(R2_full) and np.isfinite(R2_dip)) else np.nan

    return dict(label=label, A2=A2, psi2=psi2, dR2=dR2, R2_full=R2_full)


def source_errors(Smaj, Smin, Spa, bmin, flux, peak, rms):
    """
    Flux and position uncertainties. RA/Dec errors returned in degrees.
    """
    rho = sqrt(pi / (8.0 * log(2.0)) * Smaj * Smin / (bmin * bmin / 4.0) * (peak * peak) / (rms * rms))
    dA = sqrt(2.0) * (peak / rho)
    dI = sqrt(2.0) * (flux / rho)
    dx = sqrt(1.0 / (4.0 * log(2.0)) * (Smaj / rho))
    dy = sqrt(1.0 / (4.0 * log(2.0)) * (Smin / rho))
    spa_rad = Spa * pi / 180.0
    sin2 = sin(spa_rad) ** 2
    cos2 = cos(spa_rad) ** 2
    dalpha = sqrt(dx * dx * sin2 + dy * dy * cos2)
    ddelta = sqrt(dx * dx * cos2 + dy * dy * sin2)
    return (dI, dA, dalpha / 3600.0, ddelta / 3600.0)  # -> degrees


def separation(RA1, Dec1, 
               RA2, Dec2, 
               RA3, Dec3):
    return sqrt((RA1 - RA2) ** 2 + (Dec1 - Dec2) ** 2) - sqrt((RA1 - RA3) ** 2 + (Dec1 - Dec3) ** 2)


def separation_components(ra1, dec1, ra2, dec2):
    """Return east and north arcsecond separations using astropy SkyCoord.spherical_offsets_to."""
    c1 = SkyCoord(ra=ra1 * u.deg, dec=dec1 * u.deg, frame="icrs")
    c2 = SkyCoord(ra=ra2 * u.deg, dec=dec2 * u.deg, frame="icrs")
    dlon, dlat = c1.spherical_offsets_to(c2)  # east+, north+
    return dlon.to(u.arcsec).value, dlat.to(u.arcsec).value


def separation_asec(ra1, dec1, ra2, dec2):
    c1 = SkyCoord(ra=ra1 * u.deg, dec=dec1 * u.deg, frame="icrs")
    c2 = SkyCoord(ra=ra2 * u.deg, dec=dec2 * u.deg, frame="icrs")
    return c1.separation(c2).to(u.arcsec).value


def substrctr_fts(fitsname, source_table):
    """Subtract sources from a FITS image — (placeholder)."""
    pass


def pos_varCMC1xCMC2v00():
    """Deprecated legacy routine (kept for reference)."""
    pass


def _ensure_dir(d):
    os.makedirs(d, exist_ok=True)
    return d


def _docx_add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    if level == 0:
        run = p.add_run(text)
        run.bold = True
        run.font.size = None
    else:
        p.add_run(text)
    return p


def _docx_add_caption(doc, text):
    doc.add_paragraph(text)
    return


def _set_table_borders(table, width_pt=1.0, color="000000"):
    """Ensure table borders are visible with the requested width."""
    size = builtins.max(int(round(width_pt * 8)), 1)  # Word expects 1/8 pt units
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.append(tblPr)
    borders = tblPr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tblPr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = qn(f"w:{edge}")
        element = borders.find(tag)
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), str(size))
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def _save_fig(figpath):
    plt.tight_layout()
    plt.savefig(figpath, dpi=150, bbox_inches='tight')
    plt.close()


# === CIRCULAR STATS HELPERS (paste near other helpers) ===
def _circ_stats_rayleigh(theta_deg):
    """
    Circular stats + Rayleigh test for non-uniformity of angles.
    theta_deg: array-like angles in degrees (any wrap), NaNs allowed.

    Returns dict with:
      n         : sample size
      r         : mean resultant length in [0,1] (0 = uniform, 1 = all same angle)
      mean_deg  : mean direction (deg, in [0,360))
      p         : Rayleigh test p-value (small p -> non-uniform /significant clustering)
    """
    th = np.asarray(theta_deg, dtype=float)
    m  = np.isfinite(th)
    th = th[m]
    n = th.size
    if n < 5:
        return dict(n=n, r=np.nan, mean_deg=np.nan, p=np.nan)

    rad = np.deg2rad(th % 360.0)
    C = np.sum(np.cos(rad))
    S = np.sum(np.sin(rad))
    R = np.hypot(C, S)
    r = R / n
    mean = (np.rad2deg(np.arctan2(S, C)) + 360.0) % 360.0

    # Rayleigh test (approximate p-value with small-sample correction)
    # z = R^2 / n
    z = (R * R) / n
    # p ≈ exp(-z) * (1 + (2z - z^2)/(4n) - (24z - 132z^2 + 76z^3 - 9z^4)/(288 n^2))
    p = np.exp(-z) * (1.0 + (2.0*z - z*z)/(4.0*n)
                      - (24.0*z - 132.0*z*z + 76.0*z**3 - 9.0*z**4)/(288.0*n*n))
    p = float(np.clip(p, 0.0, 1.0))
    return dict(n=n, r=float(r), mean_deg=float(mean), p=p)


def _rho_bm_summary(rho_bm):
    """Return median and 84th percentile of ρ/Bmaj (ignore NaNs)."""
    x = np.asarray(rho_bm, dtype=float)
    m = np.isfinite(x)
    if not np.any(m):
        return dict(median=np.nan, p84=np.nan)
    return dict(median=float(np.nanmedian(x)), p84=float(np.nanpercentile(x, 84)))


def _stats_table_from_offsets(Dra_arcsec, Ddec_arcsec, cbmaj_arcsec, cbmin_arcsec, rho_bm_main=None):
    """Compute means/stds in arcsec and in beam units, return rows suited for a docx table."""
    rows = []
    for name, arr, scale in [
        ('ΔRA (arcsec)', Dra_arcsec, 1.0),
        ('ΔDec (arcsec)', Ddec_arcsec, 1.0),
        ('ΔRA / Bmaj', np.array(Dra_arcsec) / cbmaj_arcsec, 1.0),
        ('ΔDec / Bmaj', np.array(Ddec_arcsec) / cbmaj_arcsec, 1.0),
        ('ΔRA / Bmin', np.array(Dra_arcsec) / cbmin_arcsec, 1.0),
        ('ΔDec / Bmin', np.array(Ddec_arcsec) / cbmin_arcsec, 1.0),
    ]:
        arr = np.asarray(arr)
        rows.append([
            name,
            f"{np.nanmean(arr):.5g}",
            f"{np.nanstd(arr):.5g}",
            f"{np.nanmedian(arr):.5g}",
            f"{np.nanpercentile(arr, 16):.5g}",
            f"{np.nanpercentile(arr, 84):.5g}",
        ])
    if rho_bm_main is not None:
        arr = np.asarray(rho_bm_main)
        rows.append([
            'rho/Bmaj (main)',
            f"{np.nanmean(arr):.5g}",
            f"{np.nanstd(arr):.5g}",
            f"{np.nanmedian(arr):.5g}",
            f"{np.nanpercentile(arr, 16):.5g}",
            f"{np.nanpercentile(arr, 84):.5g}",
        ])
    return rows


def _docx_add_table(doc, rows):
    tbl = doc.add_table(rows=1, cols=6)
    hdr = tbl.rows[0].cells
    hdr[0].text = 'Quantity'
    hdr[1].text = 'Mean'
    hdr[2].text = 'Std'
    hdr[3].text = 'Median'
    hdr[4].text = 'P16'
    hdr[5].text = 'P84'
    for r in rows:
        row_cells = tbl.add_row().cells
        for i, val in enumerate(r):
            row_cells[i].text = str(val)
    _set_table_borders(tbl, width_pt=1.0)


def _legend_outside(ax):
    """Place legend outside the plotting area on the right."""
    handles, labels = ax.get_legend_handles_labels()
    if handles:
        ax.legend(handles, labels, loc='center left', bbox_to_anchor=(1.02, 0.5), borderaxespad=0., fontsize='x-small')


def _enforce_square_axes(ax):
    """Ensure plot box is square so circles stay circular."""
    try:
        ax.set_box_aspect(1)
    except AttributeError:  # fallback for old Matplotlib
        ax.set_aspect('equal', adjustable='box')


def _add_enclosing_ellipse(ax, x, y, color, label):
    """
    Draw an enclosing **axis-aligned** ellipse for the dataset (x, y):
      - Center at (mean(x), mean(y))
      - Width  = 2 * max|x - mean_x|
      - Height = 2 * max|y - mean_y|
      - Edge color = dataset color; no fill; alpha ~0.5
      - Also place an 'x' at the center with matching color
    """
    x = np.asarray(x, dtype=float)
    y = np.asarray(y, dtype=float)
    mask = np.isfinite(x) & np.isfinite(y)
    if not np.any(mask):
        return
    x = x[mask]; y = y[mask]
    xm = float(np.nanmean(x)); ym = float(np.nanmean(y))
    xme = float(np.std(x)); yme = float(np.std(y))
    rx = float(np.nanmax(np.abs(x - xm)))
    ry = float(np.nanmax(np.abs(y - ym)))
    if not np.isfinite(rx) or not np.isfinite(ry):
        return
    width  = 2.0 * rx
    height = 2.0 * ry
    ell = Ellipse((xm, ym), width=width, height=height,
                  angle=0.0, fill=False, lw=1.6, alpha=0.5,
                  edgecolor=color, label=f"{label} ellipse")
    ax.add_patch(ell)
    ax.plot(xm, ym, 'x', color=color, markersize=8, mew=2, label=f"{label}\n center [{xm:.2f}±{xme:.2f}, {ym:.2f}±{yme:.2f}]")


def pos_varCMC1xCMC2(
    ra_ref, dec_ref, x_ref, y_ref,
    ra_other, dec_other, x_other, y_other,
    ref_fits_path, other_fits_path,
    base_ref=None, base_other=None,
    xmatch_table_path=None,
):
    """
    Plot positional variations between two cross-matched catalogues measured on two images.
    Uses BMAJ/BMIN and CRPIX from the reference FITS header.

    In addition to plots, writes a DOCX report:
      - saved into 'crossmatched-positions/ref_x_<basename(other_fits)>.docx'
      - includes plots, a summary table, and physical interpretations.

    Returns a dict with figure paths, output dir, and summary arrays.
    """
    # base names for titles/labels
    if base_ref is None:
        try:
            base_ref = os.path.basename(ref_fits_path).replace('.fits', '')[:10]
        except Exception:
            base_ref = os.path.basename(ref_fits_path).replace('.fits', '')
    if base_other is None:
        try:
            base_other = os.path.basename(other_fits_path).replace('.fits', '')[:10]
        except Exception:
            base_other = os.path.basename(other_fits_path).replace('.fits', '')

    # read header
    hdr_ref = fits.getheader(ref_fits_path)
    hdr_other = fits.getheader(other_fits_path)
    if "BMAJ" not in hdr_ref or "BMIN" not in hdr_ref:
        raise ValueError("Reference FITS header missing BMAJ/BMIN.")
    cbmaj_arcsec = float(hdr_ref["BMAJ"]) * 3600.0
    cbmin_arcsec = float(hdr_ref["BMIN"]) * 3600.0
    if "CRPIX1" not in hdr_ref or "CRPIX2" not in hdr_ref:
        raise ValueError("Reference FITS header missing CRPIX1/CRPIX2.")
    if "CRPIX1" not in hdr_other or "CRPIX2" not in hdr_other:
        raise ValueError("Other FITS header missing CRPIX1/CRPIX2.")
    
    x0 = float(hdr_ref["CRPIX1"]); x2_0 = float(hdr_other["CRPIX1"])
    y0 = float(hdr_ref["CRPIX2"]); y2_0 = float(hdr_other["CRPIX2"])

    # arrays
    ra1 = np.asarray(ra_ref, dtype=float)
    dec1 = np.asarray(dec_ref, dtype=float)
    ra2 = np.asarray(ra_other, dtype=float)
    dec2 = np.asarray(dec_other, dtype=float)
    x1 = np.asarray(x_ref, dtype=float); y1 = np.asarray(y_ref, dtype=float)
    x2 = np.asarray(x_other, dtype=float); y2 = np.asarray(y_other, dtype=float)

    # pixel polar angles around CRPIX
    phi1_deg = np.degrees(np.arctan2(y1 - y0, x1 - x0)) # east of north (i.e. +x = 0°, +y = 90°) pix polar angle for ref catalogue
    phi2_deg = np.degrees(np.arctan2(y2 - y0, x2 - x0)) # east of north pix polar angle for test/cmc2 catalogue (uses ref CRPIX)

    # sky offsets & theta (uses your RA_Dec_offsets)
    Dra_arcsec, Ddec_arcsec = RA_Dec_offsets(ra1, dec1, ra2, dec2)
    theta_deg = np.degrees(np.arctan2(Ddec_arcsec, Dra_arcsec))

    # normalized variants
    Dra_bm   = Dra_arcsec  / cbmaj_arcsec
    Ddec_bm  = Ddec_arcsec / cbmaj_arcsec
    theta_bm = theta_deg   / cbmaj_arcsec  # dimensionally odd, preserved for continuity
    
    # === magnitude of the offset in beam units ===
    rho_arcsec = np.hypot(Dra_arcsec, Ddec_arcsec)
    rho_bm     = rho_arcsec / cbmaj_arcsec
    rho_bm_main = np.asarray(rho_bm)

    # === Quadrupole checks (paste after Dra_bm / Ddec_bm are computed) ===
    quad_dra_phi1 = _harmonic_quad_summary(phi1_deg, Dra_bm, label="ΔRA/Bmaj vs φ1")
    quad_dra_phi2 = _harmonic_quad_summary(phi2_deg, Dra_bm, label="ΔRA/Bmaj vs φ2")
    quad_ddec_phi1 = _harmonic_quad_summary(phi1_deg, Ddec_bm, label="ΔDec/Bmaj vs φ1")
    quad_ddec_phi2 = _harmonic_quad_summary(phi2_deg, Ddec_bm, label="ΔDec/Bmaj vs φ2")

    # Pre-format one-liners for the DOCX “Interpretation” text
    def _quad_line(q):
        if not np.isfinite(q["A2"]) or not np.isfinite(q["dR2"]):
            return f'{q["label"]}: insufficient data.'
        return (f'{q["label"]}: A₂={q["A2"]:.3g}, ψ₂={q["psi2"]:.1f}°, '
                f'ΔR²={q["dR2"]:.3g} (unique variance from 2φ).')

    quad_summary_dra  = _quad_line(quad_dra_phi1) + "  " + _quad_line(quad_dra_phi2)
    quad_summary_ddec = _quad_line(quad_ddec_phi1) + "  " + _quad_line(quad_ddec_phi2)

    print("\n---------\nQuadrupole checks summary:")
    print("[QUAD] ΔRA/Bmaj vs φ:", quad_summary_dra)
    print("[QUAD] ΔDec/Bmaj vs φ:", quad_summary_ddec)

    # === NEW: circular stats for θ and summary for ρ/Bmaj ===
    circ_theta = _circ_stats_rayleigh(theta_deg)          # r, mean, Rayleigh p
    rho_stats  = _rho_bm_summary(rho_bm)                  # median, P84

    # Pre-format one-liners for report & console
    theta_stats_line = (
        f"θ circular stats → n={circ_theta['n']}, r={circ_theta['r']:.3f}, "
        f"mean={circ_theta['mean_deg']:.1f}°, Rayleigh p={circ_theta['p']:.3g}."
    )
    rayleigh_label = (
        "strong evidence of clustering/non-uniformity" if circ_theta['p'] <= 0.01 else
        "consistent with marginal uniformity" if circ_theta['p'] <= 0.05 else
        "consistent with insignificant uniformity"
    )
    circular_stats_oneliner = "\n".join([
        f"θ circular snapshot (n={circ_theta['n']}):",
        f"  r = {circ_theta['r']:.3f} (0 ⇒ uniform, 1 ⇒ aligned)",
        f"  mean θ = {circ_theta['mean_deg']:.1f}° east of north",
        f"  Rayleigh p = {circ_theta['p']:.3g} ({rayleigh_label})",
        "  Pair this with ρ/Bmaj spread to gauge directional systematics.",
    ])
    rho_stats_line = (
        f"ρ/Bmaj stats → median={rho_stats['median']:.3g}, P84={rho_stats['p84']:.3g}."
    )

    # Console echo so results are visible on run
    print("[STATS] " + theta_stats_line)
    for line in circular_stats_oneliner.splitlines():
        print("[STATS] " + line)
    print("[STATS] " + rho_stats_line)

    # output directories + filenames
    if xmatch_table_path is not None:
        outdir = _ensure_dir(os.path.dirname(xmatch_table_path)+"/crossmatched-positions")
    else:
        outdir = _ensure_dir(os.path.dirname(other_fits_path)+"/crossmatched-positions")
    
    base_other_full = os.path.basename(xmatch_table_path).replace(".fits", "")
    report_docx = os.path.join(outdir, f"{base_ref}_x_{base_other_full}_astrometry.docx")

    fig1 = os.path.join(outdir, f"DRA-vs-phi-{base_ref}x{base_other}.png")
    fig2 = os.path.join(outdir, f"DDEC-vs-phi-{base_ref}x{base_other}.png")
    fig3 = os.path.join(outdir, f"DRA-vs-DDEC-bmaj-{base_ref}x{base_other}.png")
    fig4 = os.path.join(outdir, f"DRA-vs-DDEC-arcsec-{base_ref}x{base_other}.png")
    fig5 = os.path.join(outdir, f"theta-vs-phi-{base_ref}x{base_other}.png")
    fig6 = os.path.join(outdir, f"rho-over-bmaj-vs-phi-{base_ref}x{base_other}.png")

    # --- Generate plots (and save into outdir) ---
    # 1) ΔRA/Bmaj vs φ  [two datasets on same axes → draw ellipses]
    plt.figure(figsize=(8.5,5))
    plt.title(f'{base_ref} x {base_other} — ΔRA vs φ')
    h1, = plt.plot(Dra_bm,  phi1_deg, 'r+', label=r'$\phi_1$')
    h2, = plt.plot(Dra_bm,  phi2_deg, 'k.', label=r'$\phi_2$')
    ax = plt.gca()
    _add_enclosing_ellipse(ax, Dra_bm, phi1_deg, color=h1.get_color(), label=f"{base_ref}x{base_other} φ1")
    _add_enclosing_ellipse(ax, Dra_bm, phi2_deg, color=h2.get_color(), label=f"{base_ref}x{base_other} φ2")
    plt.xlabel('ΔRA / Bmaj'); plt.ylabel('φ (deg)')
    ax.grid(True, linestyle=':', linewidth=0.7, alpha=0.6)
    _legend_outside(ax)
    plt.tight_layout(rect=[0,0,0.78,1])  # leave room right for legend
    _save_fig(fig1)

    # 2) ΔDec/Bmaj vs φ  [two datasets → ellipses]
    plt.figure(figsize=(8.5,5))
    plt.title(f'{base_ref} x {base_other} — ΔDec vs φ')
    h1, = plt.plot(Ddec_bm, phi1_deg, 'r+', label=r'$\phi_1$')
    h2, = plt.plot(Ddec_bm, phi2_deg, 'k.', label=r'$\phi_2$')
    ax = plt.gca()
    _add_enclosing_ellipse(ax, Ddec_bm, phi1_deg, color=h1.get_color(), label=f"{base_ref}x{base_other} φ1")
    _add_enclosing_ellipse(ax, Ddec_bm, phi2_deg, color=h2.get_color(), label=f"{base_ref}x{base_other} φ2")
    plt.xlabel('ΔDec / Bmaj'); plt.ylabel('φ (deg)')
    ax.grid(True, linestyle=':', linewidth=0.7, alpha=0.6)
    _legend_outside(ax)
    plt.tight_layout(rect=[0,0,0.78,1])
    _save_fig(fig2)

    # 3) ΔRA vs ΔDec (in beam-major units)  [single dataset → no ellipse]
    plt.figure(figsize=(7,6))
    plt.title(f'{base_ref} x {base_other} — ΔRA vs ΔDec (cbmaj)')
    h3, = plt.plot(Dra_bm, Ddec_bm, 'k.', alpha=0.7, label=f'{base_ref}x{base_other}')
    plt.xlabel('ΔRA / Bmaj'); plt.ylabel('ΔDec / Bmaj'); plt.axis('equal')
    ax = plt.gca()
    _add_enclosing_ellipse(ax, Dra_bm, Ddec_bm, color=h3.get_color(), label="enclosing ellipse")
    ax.grid(True, linestyle=':', linewidth=0.7, alpha=0.6)
    _legend_outside(ax)
    plt.tight_layout(rect=[0,0,0.78,1])
    _save_fig(fig3)

    # 4) ΔRA vs ΔDec (in arcsec)  [single dataset → no ellipse]
    plt.figure(figsize=(7,6))
    plt.title(f'{base_ref} x {base_other} — ΔRA vs ΔDec (arcsec)')
    h4, = plt.plot(Dra_arcsec, Ddec_arcsec, 'k.', alpha=0.7, label=f'{base_ref}x{base_other}')
    plt.xlabel('ΔRA (arcsec)'); plt.ylabel('ΔDec (arcsec)'); plt.axis('equal')
    ax = plt.gca()
    _add_enclosing_ellipse(ax, Dra_arcsec, Ddec_arcsec, color=h4.get_color(), label="enclosing ellipse")
    ax.grid(True, linestyle=':', linewidth=0.7, alpha=0.6)
    _enforce_square_axes(ax)
    _legend_outside(ax)
    plt.tight_layout(rect=[0,0,0.78,1])
    _save_fig(fig4)

    # 5) θ/Bmaj vs φ  [two datasets → ellipses]
    plt.figure(figsize=(8.5,5))
    plt.title(f'{base_ref} x {base_other} — θ vs φ')
    h1, = plt.plot(theta_bm, phi1_deg, 'r+', label=r'$\phi_1$')
    h2, = plt.plot(theta_bm, phi2_deg, 'k.', label=r'$\phi_2$')
    ax = plt.gca()
    _add_enclosing_ellipse(ax, theta_bm, phi1_deg, color=h1.get_color(), label=f"{base_ref}x{base_other} φ1")
    _add_enclosing_ellipse(ax, theta_bm, phi2_deg, color=h2.get_color(), label=f"{base_ref}x{base_other} φ2")
    plt.xlabel('θ (deg) / Bmaj'); plt.ylabel('φ (deg)')
    ax.grid(True, linestyle=':', linewidth=0.7, alpha=0.6)
    _legend_outside(ax)
    plt.tight_layout(rect=[0,0,0.78,1])
    _save_fig(fig5)

    # 6) ρ/Bmaj vs φ  [two datasets → ellipses; physically meaningful normalization]
    plt.figure(figsize=(8.5,5))
    plt.title(f'{base_ref} x {base_other} — ρ/Bmaj vs φ')
    h1, = plt.plot(rho_bm, phi1_deg, 'r+', label=r'$\phi_1$')
    h2, = plt.plot(rho_bm, phi2_deg, 'k.', label=r'$\phi_2$')
    ax = plt.gca()
    _add_enclosing_ellipse(ax, rho_bm, phi1_deg, color=h1.get_color(), label=f"{base_ref}x{base_other} φ1")
    _add_enclosing_ellipse(ax, rho_bm, phi2_deg, color=h2.get_color(), label=f"{base_ref}x{base_other} φ2")
    plt.xlabel('ρ / Bmaj'); plt.ylabel('φ (deg)')
    ax.grid(True, linestyle=':', linewidth=0.7, alpha=0.6)
    _legend_outside(ax)
    plt.tight_layout(rect=[0,0,0.78,1])
    _save_fig(fig6)

    # --- Build summary stats (from offsets) ---
    stats_rows = _stats_table_from_offsets(Dra_arcsec, Ddec_arcsec, cbmaj_arcsec, cbmin_arcsec, rho_bm_main=rho_bm_main)

    # --- Write DOCX report ---
    doc = Document()
    _docx_add_heading(doc, f"Cross-matched positional analysis: {base_ref} x {base_other}", level=0)
    doc.add_paragraph(f"Reference image: {os.path.basename(ref_fits_path)}")
    doc.add_paragraph(f"Other image: {os.path.basename(other_fits_path)}")
    if xmatch_table_path:
        doc.add_paragraph(f"Cross-match table: {os.path.basename(xmatch_table_path)}")

    # Add plots with captions + physical interpretations
    for fp, cap, interp in [
        (
            fig1,
            "Figure 1: ΔRA/Bmaj vs pixel angle φ for reference (red, φ₁) and other (black, φ₂); ellipses show scatter and 'x' marks mean.",
            "φ is the position angle of each matched source around the image centre. If ΔRA wiggles with φ, the calibration is direction dependent "
            "(think beam squint, pointing, residual w-term). Quadrupole → "
            f"{quad_summary_dra}. We fit sin/cos terms up to 2φ; the a2/b2 pair measure how strong the four-lobed pattern is, and ΔR² tells how much "
            "extra variance those terms explain (≈0.05 is modest & ≥0.10 is clear. Bigger ΔR² ⇒ stronger, real 2φ structure. Pro tip"
            "Pro tip: Make the playing field symmetric: cross-match PB-corr ↔ PB-corr (or non-PB ↔ non-PB). Consistency beats cleverness here."
            "-- mixing a PB-corrected catalogue with a non-PB-corrected one doesn’t change the true sky positions, "
            "but it changes the way errors breathe with field angle. That mismatch can invent or amplify "
            "2φ (quadrupolar) structure in your ΔRA/ΔDec vs φ diagnostics—even if the sky is innocent.\n"
            "NB: One φ origin/ref pixel coord to diagnose direction-dependent effects relative to one field center "
            " (your reference beam, which you also use for Bmaj normalization -- i.e. φ₁ and φ₂ both use the reference image CRPIX)."
        ),
        (
            fig2,
            "Figure 2: ΔDec/Bmaj vs pixel angle φ for reference (red) and other (black); ellipses show scatter and 'x' marks mean.",
            "Same diagnostic but now for Dec offsets. Smooth φ trends reveal field-angle distortions, while small ellipses mean the behaviour is stable across "
            "the footprint. Quadrupole → "
            f"{quad_summary_ddec}, interpreted exactly as above for ΔRA.",
        ),
        (
            fig3,
            "Figure 3: ΔRA vs ΔDec in beam-major units with equal axes.",
            "Offsets plotted in synthesised-beam units. A centred, roughly circular cloud means noise-dominated astrometry; a displaced or stretched cloud "
            "flags a coherent shift or a preferred error direction.",
        ),
        (
            fig4,
            "Figure 4: ΔRA vs ΔDec in arcseconds with equal axes.",
            "Same scatter plot but in absolute arcseconds. Use this when comparing against specification numbers or external catalogues.",
        ),
        (
            fig5,
            "Figure 5: θ (deg)/Bmaj vs pixel angle φ with ellipses and means marked.",
            (
                "θ is the bearing of each offset (east of north). If the points cluster, a single dominant shift is present; if θ varies with φ, "
                "geometry-driven effects are active. Dividing by Bmaj simply keeps units consistent with the other plots. "
                f"{theta_stats_line}\n{circular_stats_oneliner}\n"
                "-- What you should hope to see (no dominant shift):\n Uniform θ distribution (no preferred angle), i.e., directionless noise. "
                "Small offset magnitudes: ρ/Bmaj well below 1 — many pipelines aim for median(ρ/Bmaj) ≪ 0.1 for well-behaved astrometry. "
                "No φ dependence: neither θ nor ρ/Bmaj changing with φ; your ellipses in the overlay plots should be compact and centred (the ‘x’ near the "
                "origin of the cloud), with similar shapes across scans.\n"
                "** When “random” isn’t entirely benign:** If θ looks uniform but ρ/Bmaj is large (or grows with φ), you’ve got isotropic but big errors—still "
                "a problem, just not directionally biased. If θ looks random yet varies with φ (e.g., different scatter width at certain φ), that hints at "
                "direction-dependent SNR/beam effects masking a weak preferred angle."
            ),
        ),
        (
            fig6,
            "Figure 6: ρ/Bmaj vs pixel angle φ for reference (red) and other (black); ellipses show scatter and 'x' marks mean.",
            (
                "ρ is the radial offset √(ΔRA²+ΔDec²). Normalising by Bmaj makes the metric dimensionless so you can read it in ‘fractions of a beam.’ If "
                "ρ/Bmaj changes with φ you are seeing direction-dependent residuals; a displaced ellipse centre highlights a bias in the mean offset "
                "strength. "
                f"{rho_stats_line}"
            ),
        ),
    ]:
        if os.path.exists(fp):
            doc.add_picture(fp, width=Inches(6.5))
            _docx_add_caption(doc, cap)
            doc.add_paragraph("Interpretation: " + interp)

    doc.add_page_break()
    _docx_add_heading(doc, "Summary statistics of positional offsets", level=1)
    _docx_add_table(doc, stats_rows)
    _docx_add_caption(doc, "Table 1: Summary metrics for ΔRA and ΔDec computed from the cross-matched catalogue. "
                      "the 16th and 84th percentiles are a robust, model-light way to say “about one sigma either "
                      "side of the median” if the distribution is roughly symmetric. They’re handy because they "
                      "don’t assume Gaussian noise but still map to familiar intuition.")

    doc.save(report_docx)

    return {
        'cbmaj_arcsec': cbmaj_arcsec,
        'cbmin_arcsec': cbmin_arcsec,
        'x0': x0, 'y0': y0, 'x2_0': x2_0, 'y2_0': y2_0,
        'rho_arcsec': rho_arcsec, 'rho_bm': rho_bm,
        'Dra_arcsec': Dra_arcsec, 'Ddec_arcsec': Ddec_arcsec, 'theta_deg': theta_deg,
        'Dra_bm': Dra_bm, 'Ddec_bm': Ddec_bm, 'theta_bm': theta_bm,
        'rho_bm_main': rho_bm_main,
        'phi1_deg': phi1_deg, 'phi2_deg': phi2_deg,
        'figs': [fig1, fig2, fig3, fig4, fig5, fig6],
        'report_docx': report_docx,
        'outdir': outdir,
        'base_ref': base_ref,
        'base_other': base_other,
    }


def get_SUMSSXcmcX(tablename, otherfits, reffits):
    """
    Example routine for SUMSS cross-match — left intact; not used by the new wrapper.
    """
    t = Table.read(tablename)
    hdr_ref = fits.getheader(reffits)
    cbmaj_arcsec = float(hdr_ref["BMAJ"]) * 3600.0
    ra_cmcX  = t['RAJ2000'];  era_cmcX  = t['e_RAJ2000']
    dec_cmcX = t['DEJ2000'];  edec_cmcX = t['e_DEJ2000']
    ra_SM    = t['_RAJ2000']; era_SM    = t['e_RAJ2000']
    dec_SM   = t['_DEJ2000']; edec_SM   = t['e_DEJ2000']
    xpx_cmcX = t['Xposn'];     ypx_cmcX  = t['Yposn']
    xpx_SM   = t['Xpos'];      ypx_SM    = t['Ypos']

    # Convert to plain deg values for offsets
    ra1  = u.Quantity(ra_cmcX,  u.deg, copy=False).to_value(u.deg)
    dec1 = u.Quantity(dec_cmcX, u.deg, copy=False).to_value(u.deg)
    ra2  = u.Quantity(ra_SM,    u.deg, copy=False).to_value(u.deg)
    dec2 = u.Quantity(dec_SM,   u.deg, copy=False).to_value(u.deg)

    Dra, Ddec = RA_Dec_offsets(ra1, dec1, ra2, dec2)  # arcsec

    print("> Position errors:")
    print(f" - mean(Delta_RA) : {np.nanmean(Dra)/cbmaj_arcsec:.3g}, sd(Delta_RA) : {np.nanstd(Dra)/cbmaj_arcsec:.3g} [cbmaj]")
    print(f" - mean(Delta_Dec): {np.nanmean(Ddec)/cbmaj_arcsec:.3g}, sd(Delta_Dec): {np.nanstd(Ddec)/cbmaj_arcsec:.3g} [cbmaj]")


def RA_Dec_offsets00(ra1, dec1, ra2, dec2):
    """Legacy RA/Dec offsets in arcsec (east+, north+)."""
    Dra  = (ra2 - ra1) * 3600.0 * np.cos(np.radians(0.5 * (dec1 + dec2)))
    Ddec = (dec2 - dec1) * 3600.0
    return Dra, Ddec


def RA_Dec_offsets(ra1, dec1, ra2, dec2):
    """
    Accurate RA/Dec offsets (east+, north+) in arcseconds using SkyCoord.spherical_offsets_to.
    """
    c1 = SkyCoord(ra=ra1 * u.deg, dec=dec1 * u.deg, frame="icrs")
    c2 = SkyCoord(ra=ra2 * u.deg, dec=dec2 * u.deg, frame="icrs")

    dlon, dlat = c1.spherical_offsets_to(c2)  # east+, north+
    return dlon.to(u.arcsec).value, dlat.to(u.arcsec).value


def get_cmc2Xcmc1(tablename):
    """
    Astrometry for sources cross-matched between CMC1 and CMC2.
    Returns RAs/Decs (as read), errors, arcsec offsets, pixel positions, and theta (deg).
    """
    t = Table.read(tablename)

    ra_cmc1  = t['RA_1'];  era_cmc1  = t['E_RA_1']
    ra_cmc2  = t['RA_2'];  era_cmc2  = t['E_RA_2']
    dec_cmc1 = t['DEC_1']; edec_cmc1 = t['E_DEC_1']
    dec_cmc2 = t['DEC_2']; edec_cmc2 = t['E_DEC_2']
    xpx_cmc1 = t['Xposn_1']; ypx_cmc1 = t['Yposn_1']
    xpx_cmc2 = t['Xposn_2']; ypx_cmc2 = t['Yposn_2']

    # Convert to plain degree floats for math
    ra1  = u.Quantity(ra_cmc1,  u.deg, copy=False).to_value(u.deg)
    dec1 = u.Quantity(dec_cmc1, u.deg, copy=False).to_value(u.deg)
    ra2  = u.Quantity(ra_cmc2,  u.deg, copy=False).to_value(u.deg)
    dec2 = u.Quantity(dec_cmc2, u.deg, copy=False).to_value(u.deg)

    # theta in degrees (angle of vector from CMC2->CMC1)
    theta = np.degrees(np.arctan2(dec1 - dec2, ra1 - ra2))

    # arcsec offsets (east+, north+)
    Dra, Ddec = RA_Dec_offsets(ra1, dec1, ra2, dec2)

    return (ra_cmc1, era_cmc1, ra_cmc2, era_cmc2,
            dec_cmc1, edec_cmc1, dec_cmc2, edec_cmc2,
            Dra, Ddec, xpx_cmc1, ypx_cmc1, xpx_cmc2, ypx_cmc2, theta)


def CrossMatch(RA1_arr, DEC1_arr, RA2_arr, DEC2_arr):
    """
    Cross-match two coordinate sets using SkyCoord.match_to_catalog_sky().

    Returns
    -------
    idx : array
        Indices into RA2/DEC2 that are nearest to each RA1/DEC1.
    sep2d : Angle
        On-sky separations.
    dist3d : Quantity
        3D distances (ignore for pure sky matches).
    RA2_match, DEC2_match : array
        Matched coordinates from set 2.
    """
    c1 = SkyCoord(ra=RA1_arr * u.deg, dec=DEC1_arr * u.deg, frame='icrs')
    c2 = SkyCoord(ra=RA2_arr * u.deg, dec=DEC2_arr * u.deg, frame='icrs')

    idx, sep2d, dist3d = c1.match_to_catalog_sky(c2)
    RA2_match  = RA2_arr[idx]
    DEC2_match = DEC2_arr[idx]
    return idx, sep2d, dist3d, RA2_match, DEC2_match


def get_LinPol_PA(tablename, qname, uname):
    """Example placeholder."""
    pass


def get_LinPol_PA2(tablename, qname, uname):
    """Example placeholder for linear polarization calculation."""
    t = Table.read(tablename)
    q = t[qname]
    u = t[uname]
    P = []
    PA = []
    for i in range(len(q)):
        if np.isfinite(q[i]) and np.isfinite(u[i]):
            p  = np.hypot(q, u)
            pa = 0.5 * np.degrees(np.arctan2(u, q))
        else:
            p = np.nan
            pa = np.nan
        print(f'P: {p}, PA: {pa} +++ Q: {q}, U:{u}')
        P.append(p)
        PA.append(pa)

    P  = np.array(P)
    PA = np.array(PA)
    print(f'avg P: {np.nanmean(P)}, rms P: {np.nanstd(P)}, avgPA: {np.nanmean(PA)}')
    return (P, PA)


def analyze_ref_vs_other_with_optional_scans(
    xmatch_table_path,
    ref_fits_path,
    other_fits_path,
    baseref0=None,
    baseother0=None,
    per_scan_glob=None,
):
    """
    Streamlined wrapper:
      1) Reads the cross-match *table* between reference and other image,
      2) Calls get_cmc2Xcmc1(...) to extract positions and offsets,
      3) Calls pos_varCMC1xCMC2(...) to generate the standard five plots + DOCX,
      4) Optionally overlays *per-scan* cross-match results (matched to the same reference)
         on top of the same axes from pos_varCMC1xCMC2, with color-matched ellipses.

    Returns dict with output paths and figure lists.
    """

    if baseref0 is None:
        baseref0 = "ref"
    if baseother0 is None:
        baseother0 = os.path.basename(xmatch_table_path).replace('.fits','')

    # --- main pair ---
    (ra_ref, era_ref, ra_other, era_other,
     dec_ref, edec_ref, dec_other, edec_other,
     Dra_arcsec, Ddec_arcsec, xpx_ref, ypx_ref, xpx_other, ypx_other, theta_deg
    ) = get_cmc2Xcmc1(xmatch_table_path)

    # Produce standard outputs + get outdir
    main_res = pos_varCMC1xCMC2(
        ra_ref, dec_ref, xpx_ref, ypx_ref,
        ra_other, dec_other, xpx_other, ypx_other,
        ref_fits_path=ref_fits_path,
        other_fits_path=other_fits_path,
        xmatch_table_path=xmatch_table_path,
        base_ref=baseref0,
        base_other=baseother0,
    )
    outdir      = main_res.get('outdir', os.path.join(os.path.dirname(other_fits_path), 'crossmatched-positions'))
    base_ref    = main_res.get('base_ref', os.path.basename(ref_fits_path).replace('.fits',''))
    base_other  = main_res.get('base_other', os.path.basename(other_fits_path).replace('.fits',''))

    # If no per-scan, we're done
    if not per_scan_glob:
        return main_res

    # --- overlay per-scan tables ---
    scan_files = sorted(glob.glob(per_scan_glob))
    if not scan_files:
        print(f"[WARN] No per-scan files matched the pattern: {per_scan_glob}")
        return main_res

    # read reference header once
    hdr_ref = fits.getheader(ref_fits_path)
    hdr_other = fits.getheader(other_fits_path)
    cbmaj_arcsec = float(hdr_ref['BMAJ']) * 3600.0
    cbmin_arcsec = float(hdr_ref['BMIN']) * 3600.0

    x0 = float(hdr_ref["CRPIX1"]); x2_0 = float(hdr_other["CRPIX1"])
    y0 = float(hdr_ref["CRPIX2"]); y2_0 = float(hdr_other["CRPIX2"])

    # Prepare phi for the main pair
    phi1_deg = np.degrees(np.arctan2(ypx_ref - y0, xpx_ref - x0))
    phi2_deg = np.degrees(np.arctan2(ypx_other - y0, xpx_other - x0))

    # Normalize main pair offsets by beam major
    Dra_bm_main  = np.array(Dra_arcsec)  / cbmaj_arcsec
    Ddec_bm_main = np.array(Ddec_arcsec) / cbmaj_arcsec
    theta_bm_main= np.array(theta_deg)   / cbmaj_arcsec
    
    # main scan ρ/Bmaj
    rho_bm_main = np.hypot(Dra_bm_main, Ddec_bm_main)

    # Collect per-scan data
    overlays = []
    for i, scan_path in enumerate(scan_files, 1):
        try:
            (ra_r, era_r, ra_s, era_s,
             dec_r, edec_r, dec_s, edec_s,
             Dra_s, Ddec_s, xpx_r, ypx_r, xpx_s, ypx_s, theta_s
            ) = get_cmc2Xcmc1(scan_path)

            phi1_s = np.degrees(np.arctan2(np.array(ypx_r) - y0, np.array(xpx_r) - x0))
            phi2_s = np.degrees(np.arctan2(np.array(ypx_s) - y0, np.array(xpx_s) - x0))

            overlays.append(dict(
                label=os.path.basename(scan_path).replace('.fits',''),
                Dra_bm = np.array(Dra_s)  / cbmaj_arcsec,
                Ddec_bm= np.array(Ddec_s) / cbmaj_arcsec,
                Dra    = np.array(Dra_s),
                Ddec   = np.array(Ddec_s),
                theta  = np.array(theta_s),
                phi1   = phi1_s,
                phi2   = phi2_s,
            ))
        except Exception as e:
            print(f"[WARN] Skipping per-scan file {scan_path}: {e}")
            continue

    if not overlays:
        return main_res

    # --- make overlay figures ---
    def _save_overlay(figpath):
        plt.tight_layout()
        plt.savefig(figpath, dpi=150, bbox_inches='tight')
        plt.close()

    fig_paths = {}

    # 1) ΔRA (cbmaj) vs φ  [multiple datasets → ellipses, color-matched]
    plt.figure(figsize=(15,10))
    plt.title(f"{base_ref} x {base_other} — ΔRA vs φ (per-scan overlays)")
    h_main1, = plt.plot(Dra_bm_main, phi1_deg, '.', alpha=0.6, label='main φ1')
    h_main2, = plt.plot(Dra_bm_main, phi2_deg, '+', alpha=0.6, label='main φ2')
    for ov in overlays:
        h1, = plt.plot(ov['Dra_bm'], ov['phi1'], '.', alpha=0.8, label=ov['label']+" φ1")
        h2, = plt.plot(ov['Dra_bm'], ov['phi2'], '+', alpha=0.8, label=ov['label']+" φ2")
        ax = plt.gca()
        _add_enclosing_ellipse(ax, ov['Dra_bm'], ov['phi1'], color=h1.get_color(), label=ov['label']+" φ1")
        _add_enclosing_ellipse(ax, ov['Dra_bm'], ov['phi2'], color=h2.get_color(), label=ov['label']+" φ2")
    ax = plt.gca()
    _add_enclosing_ellipse(ax, Dra_bm_main, phi1_deg, color=h_main1.get_color(), label='main φ1')
    _add_enclosing_ellipse(ax, Dra_bm_main, phi2_deg, color=h_main2.get_color(), label='main φ2')
    plt.xlabel('ΔRA / Bmaj'); plt.ylabel('φ (deg)')
    ax.grid(True, linestyle=':', linewidth=0.7, alpha=0.6)
    _legend_outside(ax)
    plt.tight_layout(rect=[0,0,0.78,1])
    fp = os.path.join(outdir, f"DRA-vs-phi-{base_ref}x{base_other}-overlays.png"); _save_overlay(fp)
    fig_paths['dra_phi_overlays'] = fp

    # 2) ΔDec (cbmaj) vs φ
    plt.figure(figsize=(15,10))
    plt.title(f"{base_ref} x {base_other} — ΔDec vs φ (per-scan overlays)")
    h_main1, = plt.plot(Ddec_bm_main, phi1_deg, '.', alpha=0.6, label='main φ1')
    h_main2, = plt.plot(Ddec_bm_main, phi2_deg, '+', alpha=0.6, label='main φ2')
    for ov in overlays:
        h1, = plt.plot(ov['Ddec_bm'], ov['phi1'], '.', alpha=0.8, label=ov['label']+" φ1")
        h2, = plt.plot(ov['Ddec_bm'], ov['phi2'], '+', alpha=0.8, label=ov['label']+" φ2")
        ax = plt.gca()
        _add_enclosing_ellipse(ax, ov['Ddec_bm'], ov['phi1'], color=h1.get_color(), label=ov['label']+" φ1")
        _add_enclosing_ellipse(ax, ov['Ddec_bm'], ov['phi2'], color=h2.get_color(), label=ov['label']+" φ2")
    ax = plt.gca()
    _add_enclosing_ellipse(ax, Ddec_bm_main, phi1_deg, color=h_main1.get_color(), label='main φ1')
    _add_enclosing_ellipse(ax, Ddec_bm_main, phi2_deg, color=h_main2.get_color(), label='main φ2')
    plt.xlabel('ΔDec / Bmaj'); plt.ylabel('φ (deg)')
    ax.grid(True, linestyle=':', linewidth=0.7, alpha=0.6)
    _legend_outside(ax)
    plt.tight_layout(rect=[0,0,0.78,1])
    fp = os.path.join(outdir, f"DDEC-vs-phi-{base_ref}x{base_other}-overlays.png"); _save_overlay(fp)
    fig_paths['ddec_phi_overlays'] = fp

    # 3) ΔRA vs ΔDec (cbmaj units)
    plt.figure(figsize=(8.8,6.2))
    plt.title(f"{base_ref} x {base_other} — ΔRA vs ΔDec (cbmaj) per-scan overlays")
    hmain, = plt.plot(Dra_bm_main, Ddec_bm_main, 'o', alpha=0.5, label='main')
    for ov in overlays:
        h, = plt.plot(ov['Dra_bm'], ov['Ddec_bm'], '.', alpha=0.8, label=ov['label'])
        ax = plt.gca()
        _add_enclosing_ellipse(ax, ov['Dra_bm'], ov['Ddec_bm'], color=h.get_color(), label=ov['label'])
    ax = plt.gca()
    _add_enclosing_ellipse(ax, Dra_bm_main, Ddec_bm_main, color=hmain.get_color(), label='main')
    plt.xlabel('ΔRA / Bmaj'); plt.ylabel('ΔDec / Bmaj'); plt.axis('equal')
    ax.grid(True, linestyle=':', linewidth=0.7, alpha=0.6)
    _legend_outside(ax)
    plt.tight_layout(rect=[0,0,0.78,1])
    fp = os.path.join(outdir, f"DRA-vs-DDEC-bmaj-{base_ref}x{base_other}-overlays.png"); _save_overlay(fp)
    fig_paths['dra_ddec_bmaj_overlays'] = fp

    # 4) ΔRA vs ΔDec (arcsec)
    plt.figure(figsize=(8.8,6.2))
    plt.title(f"{base_ref} x {base_other} — ΔRA vs ΔDec (arcsec) per-scan overlays")
    hmain, = plt.plot(Dra_arcsec, Ddec_arcsec, 'o', alpha=0.5, label='main')
    for ov in overlays:
        h, = plt.plot(ov['Dra'], ov['Ddec'], '.', alpha=0.8, label=ov['label'])
        ax = plt.gca()
        _add_enclosing_ellipse(ax, ov['Dra'], ov['Ddec'], color=h.get_color(), label=ov['label'])
    ax = plt.gca()
    _add_enclosing_ellipse(ax, Dra_arcsec, Ddec_arcsec, color=hmain.get_color(), label='main')
    plt.xlabel('ΔRA (arcsec)'); plt.ylabel('ΔDec (arcsec)'); plt.axis('equal')
    ax.grid(True, linestyle=':', linewidth=0.7, alpha=0.6)
    _enforce_square_axes(ax)
    _legend_outside(ax)
    plt.tight_layout(rect=[0,0,0.78,1])
    fp = os.path.join(outdir, f"DRA-vs-DDEC-arcsec-{base_ref}x{base_other}-overlays.png"); _save_overlay(fp)
    fig_paths['dra_ddec_arcsec_overlays'] = fp

    # 5) θ vs φ (normalized)
    plt.figure(figsize=(15,10))
    plt.title(f"{base_ref} x {base_other} — θ vs φ (per-scan overlays)")
    h_main1, = plt.plot(theta_bm_main, phi1_deg, '.', alpha=0.6, label='main φ1')
    h_main2, = plt.plot(theta_bm_main, phi2_deg, '+', alpha=0.6, label='main φ2')
    for ov in overlays:
        theta_bm = ov['theta'] / cbmaj_arcsec
        h1, = plt.plot(theta_bm, ov['phi1'], '.', alpha=0.8, label=ov['label']+" φ1")
        h2, = plt.plot(theta_bm, ov['phi2'], '+', alpha=0.8, label=ov['label']+" φ2")
        ax = plt.gca()
        _add_enclosing_ellipse(ax, theta_bm, ov['phi1'], color=h1.get_color(), label=ov['label']+" φ1")
        _add_enclosing_ellipse(ax, theta_bm, ov['phi2'], color=h2.get_color(), label=ov['label']+" φ2")
    ax = plt.gca()
    _add_enclosing_ellipse(ax, theta_bm_main, phi1_deg, color=h_main1.get_color(), label='main φ1')
    _add_enclosing_ellipse(ax, theta_bm_main, phi2_deg, color=h_main2.get_color(), label='main φ2')
    plt.xlabel('θ (deg) / Bmaj'); plt.ylabel('φ (deg)')
    ax.grid(True, linestyle=':', linewidth=0.7, alpha=0.6)
    _legend_outside(ax)
    plt.tight_layout(rect=[0,0,0.78,1])
    fp = os.path.join(outdir, f"theta-vs-phi-{base_ref}x{base_other}-overlays.png"); _save_overlay(fp)
    fig_paths['theta_phi_overlays'] = fp

    # 6) ρ/Bmaj vs φ (per-scan overlays)
    plt.figure(figsize=(15,10))
    plt.title(f"{base_ref} x {base_other} — ρ/Bmaj vs φ (per-scan overlays)")
    h_main1, = plt.plot(rho_bm_main, phi1_deg, '.', alpha=0.6, label='main φ1')
    h_main2, = plt.plot(rho_bm_main, phi2_deg, '+', alpha=0.6, label='main φ2')

    ax = plt.gca()
    for ov in overlays:
        rho_bm = np.hypot(ov['Dra_bm'], ov['Ddec_bm'])
        h1, = plt.plot(rho_bm, ov['phi1'], '.', alpha=0.8, label=ov['label']+" φ1")
        h2, = plt.plot(rho_bm, ov['phi2'], '+', alpha=0.8, label=ov['label']+" φ2")
        _add_enclosing_ellipse(ax, rho_bm, ov['phi1'], color=h1.get_color(), label=ov['label']+" φ1")
        _add_enclosing_ellipse(ax, rho_bm, ov['phi2'], color=h2.get_color(), label=ov['label']+" φ2")

    _add_enclosing_ellipse(ax, rho_bm_main, phi1_deg, color=h_main1.get_color(), label='main φ1')
    _add_enclosing_ellipse(ax, rho_bm_main, phi2_deg, color=h_main2.get_color(), label='main φ2')
    plt.xlabel('ρ / Bmaj'); plt.ylabel('φ (deg)')
    ax.grid(True, linestyle=':', linewidth=0.7, alpha=0.6)
    _legend_outside(ax)
    plt.tight_layout(rect=[0,0,0.78,1])
    fp = os.path.join(outdir, f"rho-over-bmaj-vs-phi-{base_ref}x{base_other}-overlays.png")
    _save_overlay(fp)
    fig_paths['rho_phi_overlays'] = fp


    # Add overlays to the docx, if present, and include interpretations
    try:
        report_docx = main_res.get('report_docx')
        if report_docx and os.path.exists(report_docx):
            doc = Document(report_docx)
            _docx_add_heading(doc, "Per-scan overlays", level=1)
            overlay_entries = [
                ('dra_phi_overlays',
                 "ΔRA / Bmaj vs φ with per-scan overlays. Color-matched ellipses enclose each dataset; 'x' marks the mean.",
                 "Interpretation: Compare per-scan behavior of ΔRA across φ. Scans with larger horizontal ellipse width show more ΔRA spread; "
                 "vertical extent reflects φ coverage/scatter. Differences between scan ellipses can indicate time-variable direction-dependent errors."),
                ('ddec_phi_overlays',
                 "ΔDec / Bmaj vs φ with per-scan overlays. Color-matched ellipses enclose each dataset; 'x' marks the mean.",
                 "Interpretation: As above for Dec. Strong scan-to-scan changes in ellipse center or width suggest varying calibrations, ionosphere, or beam coupling."),
                ('dra_ddec_bmaj_overlays',
                 "ΔRA vs ΔDec (cbmaj) with per-scan overlays.",
                 "Interpretation: The 2D error cloud per scan in resolution units. Off-centered means imply per-scan registration shifts; anisotropic ellipses point to directional systematics."),
                ('dra_ddec_arcsec_overlays',
                 "ΔRA vs ΔDec (arcsec) with per-scan overlays.",
                 "Interpretation: Same as cbmaj but in arcsec for absolute context; helpful when comparing to external astrometric references."),
                ('theta_phi_overlays',
                 "θ/Bmaj vs φ with per-scan overlays. (Note: θ/Bmaj is not dimensionless; retained for continuity.)",
                 "Interpretation: Per-scan preferred offset directions. Clustering of means at a common θ suggests a stable direction of error; φ dependence signals field geometry effects."),
                ('rho_phi_overlays',
                 "ρ/Bmaj vs φ with per-scan overlays. Color-matched ellipses enclose each dataset; 'x' marks the mean.",
                 "Interpretation: Compares the magnitude of per-scan offsets (in units of beam size) around the field. "
                 "Larger ellipses or shifted centers indicate scans with systematically higher residuals or φ-dependent behavior, pointing to time-variable DD effects or calibration drift."),
            ]

            for fig_number, (path_key, cap, interp) in enumerate(overlay_entries, start=7):
                fig_path = fig_paths.get(path_key)
                if fig_path and os.path.exists(fig_path):
                    doc.add_picture(fig_path, width=Inches(6.5))
                    _docx_add_caption(doc, f"Figure {fig_number}: {cap}")
                    doc.add_paragraph(interp)
            doc.save(report_docx)
    except Exception as e:
        print(f"[WARN] Could not append overlays to DOCX: {e}")

    main_res['overlay_figures'] = fig_paths
    return main_res


def _build_argparser():
    p = argparse.ArgumentParser(description='Cross-matched positional analysis with optional per-scan overlays.')
    p.add_argument('--xmatch-table', required=True, help='FITS table path: reference x other cross-match')
    p.add_argument('--ref-fits', required=True, help='Reference image FITS path')
    p.add_argument('--other-fits', required=True, help='Other/test image FITS path')
    p.add_argument('--per-scan-glob', default=None, help='Glob pattern for per-scan cross-match FITS tables (e.g., "/data/xmatches/refXscan*.fits")')
    p.add_argument('--otherdatatag', default=None, help='string tag to identify the "other" dataset in filenames (for labeling); defaults to basename of other-fits without .fits')
    return p


def _cli():
    parser = _build_argparser()
    args = parser.parse_args()
    res = analyze_ref_vs_other_with_optional_scans(
        xmatch_table_path=args.xmatch_table,
        ref_fits_path=args.ref_fits,
        other_fits_path=args.other_fits,
        baseref0=None,
        baseother0=args.otherdatatag,
        per_scan_glob=args.per_scan_glob,
    )
    outdir = res.get('outdir')
    report = res.get('report_docx')
    print(f"\n[DONE] Outputs in:\n\t {outdir}")
    if report:
        print(f"[DOCX]:\n {report}\n")


if __name__ == '__main__':
    _cli()
